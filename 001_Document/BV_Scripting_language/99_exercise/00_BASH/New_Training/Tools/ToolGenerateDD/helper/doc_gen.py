#!/usr/bin/python

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import xmlparser
import sys
import template_str
import os
import subprocess
import string
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import ntpath
import json
import re

def to_asscii(text):
    if text is None:
        return 'None'
    else:
        return text


def width_percent(percent):
    return percent * 5275580 #unit is EMU  (1 inch = 914400EMUs)
    

def format_column_width(table, widths):
    if len(table.columns) > len(widths):
        return
    for i in range(len(table.columns)):
        column = table.columns[i]
        for c in column.cells:
            c.width = width_percent(widths[i])


def color_fill_cell(cell, text = None):
    shading_elm = parse_xml(r'<w:shd {} w:fill="d3d3d3"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    if text is not None:
        cell.text = text


def add_filed_document(doc, field):
    global field_method_count, class_count
    global class_field_tsdd
    doc.add_heading(field.Name, template_str.heading['fileds'])
    field_method_count += 1
    doc.add_paragraph(template_str.class_disigned_id %(class_count, field_method_count), style=doc.styles['DesignID'])
    table = doc.add_table(rows=4, cols=2, style='Table V.Header')
     
    class_field_tsdd[field.Name] = (template_str.class_disigned_id %(class_count, field_method_count))[:-1]
    rows = (
        ('Variable Name', field.Name),
        ('Type', field.Type),
        ('Range', field.Range),
        ('Description', field.Description),
    )
    for i in range(len(rows)):
        (title, info) = rows[i]
        row = table.rows[i]
        row.cells[0].text = to_asscii(title)
        row.cells[1].text = to_asscii(info)

    format_column_width(table, (0.3, 0.7))
    doc.add_paragraph()

field_method_count = 0
class_count = 0
REF_DB = {}
tsdd = {}
field_tsdd = {}
def add_method_document(doc, method):
    global field_method_count, class_count
    global class_tsdd
    
    field_method_count += 1
    doc.add_heading(method.Name, template_str.heading['methods'])
    doc.add_paragraph(template_str.class_disigned_id %(class_count, field_method_count), style=doc.styles['DesignID'])
    table = doc.add_table(rows=1, cols=4, style='Table V.Header')
    
    table.autofit = False
    format_column_width(table, (0.25, 0.3, 0.15, 0.3))
    
    class_tsdd[method.Name] = (template_str.class_disigned_id %(class_count, field_method_count))[:-1]
    # method name
    
    row = table.rows[0]
    heading_cell = row.cells[0]
    info_cell = row.cells[1]
    for i in range(2, len(row.cells)): info_cell.merge(row.cells[i])
    heading_cell.text = 'Method Name'
    info_cell.text = method.Name

    # arguments
    heading = []
    base_row = None
    if len(method.Args) == 0:
        #method.Args.append(("None", "None", "None"))
        row1 = table.add_row();
        row1.cells[0].text = to_asscii('Arguments')
        row1.cells[1].text = to_asscii('None')
        row1.cells[1].merge(row1.cells[3])
        
    for i in range(len(method.Args)):
        row1 = table.add_row()
        row2 = table.add_row()

        if base_row is None:
            row1.cells[0].merge(row2.cells[0])
            base_row = row1
        else:
            base_row.cells[0].merge(row1.cells[0])
            base_row.cells[0].merge(row2.cells[0])

        row1.cells[1].merge(row2.cells[1])
        row1.cells[2].merge(row1.cells[3])
        row1.cells[0].text= to_asscii('Arguments')
        row1.cells[1].text = to_asscii(method.Args[i][0]) if len(method.Args[i]) > 0 else "None"
        row1.cells[2].text = to_asscii(method.Args[i][1]) if len(method.Args[i]) > 1 else "None"
        color_fill_cell(row2.cells[2], 'Range:')
        row2.cells[3].text = to_asscii(method.Args[i][2]) if len(method.Args[i]) > 2 else "None"

    for x in range(1, len(heading)): heading[0].merge(heading[i])
    if len(heading) > 0: color_fill_cell(heading[0], 'Arguments')

    # Return value
    
    return_row1 = table.add_row()
        
    #return_row1.cells[2].merge(return_row1.cells[3])
    return_row1.cells[0].text = 'Return'
    return_row1.cells[1].text = to_asscii(method.ReturnStr[0])
    
    if to_asscii(method.ReturnStr[0]).strip() == 'None' or to_asscii(method.ReturnStr[0]).strip() == '':
        return_row1.cells[1].merge(return_row1.cells[3])
    else:
        return_row1.cells[2].merge(return_row1.cells[3])
        return_row1.cells[2].text = 'None'
        return_row2 = table.add_row()
        return_row1.cells[0].merge(return_row2.cells[0])
        return_row1.cells[1].merge(return_row2.cells[1])
        color_fill_cell(return_row2.cells[2], 'Range:')
        return_row2.cells[3].text = to_asscii(method.ReturnStr[1])
        if return_row2.cells[3].text.strip() == 'None':
            format_column_width(table, (0.25, 0.5, 0.125, 0.125))

    # Generated Value
    row = table.add_row()
    heading_cell = row.cells[0]
    info_cell = row.cells[1]
    for i in range(2, len(row.cells)): info_cell.merge(row.cells[i])
    heading_cell.text = 'Generated Value'
    info_cell.text = to_asscii(method.GeneatedValue)
    if info_cell.text.strip() == '':
        info_cell.text = 'None'

    # Description
    row = table.add_row()
    heading_cell = row.cells[0]
    info_cell = row.cells[1]
    for i in range(2, len(row.cells)): info_cell.merge(row.cells[i])
    heading_cell.text ='Description'
    
    algorithm = get_algorithm(method)
    if method.Description is not None:
        desc = to_asscii(method.Description.strip())
    else:
        desc = ''
    info_cell.text = ''
    info_cell.paragraphs[0].text = desc
    info_cell.paragraphs[0].style = 'SourceCode'
    
    if algorithm.strip() != '' and algorithm.strip() != "Algorithm:" and algorithm.strip() != "None":
        info_cell.add_paragraph('Algorithm:\n' + to_asscii(algorithm).strip(), style="SourceCode")
        
    ref = get_ref(method.Name, template_str.class_ref_req, desc)
    doc.add_paragraph(ref, style=doc.styles['RefReq'])
    doc.add_paragraph("")

def get_algorithm(method):
    with open(method.File, 'r') as f:
        lines = [line for line in f]
        i = method.Line
        algorithm = ''
        state = 0 
        while True:
            if state == 0:
                # start
                if '</algorithm>' in lines[i]:
                    state = 1
                elif '}' in lines[i]:
                    # not found
                    break
            elif state == 1:
                # get
                if '<algorithm>' in lines[i]:
                    break
                else:
                    lines[i] = lines[i].strip().replace('///', '')
                    algorithm = lines[i] + '\n' + algorithm
            i -= 1
        
    return algorithm
    
    
def get_ref(method, default, description):
    # Set ref to N/A for constructor
    if "constructor" in description or "initialize" in description:
        return template_str.class_ref_NA
    
    # Check method is error list or configuration
    method_error_list = re.search(r'(Check+\w+\d)', method)
    if method_error_list:
        name_split = re.findall(r'(Check+\w+\d)', method)
        return template_str.class_ref_errList %(name_split[0][5:8], name_split[0][11:14])
    else:
        # Get Id by configuration file
        for (key, value) in REF_DB.items():
            if method.endswith(value) or value in description:
                if key[len(key)-3:len(key)] == "001":
                    return template_str.class_ref_configBase %(key[0:len(key)-4], key)
                else:
                    return template_str.class_ref_config %key

    return default;

def add_list_fields(doc, name, class_info):
    doc.add_paragraph()
    doc.add_paragraph(name, style=doc.styles['Caption Table'])

    table = doc.add_table(rows=1, cols=4, style='Table H.Header wi.No')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # row 1
    table.rows[0].cells[0].text = 'No.'
    table.rows[0].cells[1].text = 'Name'
    table.rows[0].cells[2].text = 'Type'
    table.rows[0].cells[3].text = 'Description'
    
    empty = 1
    
    #public fields
    count = 0
    fields = class_info.PublicFields
    for i in range(len(fields)):
        count += 1
        mt = fields[i]
        row = table.add_row()
        row.cells[0].text = to_asscii('(%i)' % (count))
        row.cells[1].text = to_asscii(mt.Name)
        row.cells[2].text = "Public"
        try:
            row.cells[3].text = to_asscii(mt.Description.lstrip())
        except:
            row.cells[3].text = 'None'
        empty = 0
        
    #protected fields
    fields = class_info.ProtectedFields
    for i in range(len(fields)):
        count += 1
        mt = fields[i]
        row = table.add_row()
        row.cells[0].text = to_asscii('(%i)' % (count))
        row.cells[1].text = to_asscii(mt.Name)
        row.cells[2].text = "Protected"
        try:
            row.cells[3].text = to_asscii(mt.Description.lstrip())
        except:
            row.cells[3].text = 'None'
        empty = 0
        
    #private fields
    fields = class_info.PrivateFields
    for i in range(len(fields)):
        count += 1
        mt = fields[i]
        row = table.add_row()
        row.cells[0].text = to_asscii('(%i)' % (count))
        row.cells[1].text = to_asscii(mt.Name)
        row.cells[2].text = "Private"
        try:
            row.cells[3].text = to_asscii(mt.Description.lstrip())
        except:
            row.cells[3].text = 'None'
        empty = 0
        
    if empty == 1:
        table.add_row()
        table.rows[1].cells[0].text = '-'
        table.rows[1].cells[1].text = '-'
        table.rows[1].cells[2].text = '-'
        table.rows[1].cells[3].text = '-'

    #format_column_width(table, (0.1, 0.4, 0.5))
    table.autofit = False
    format_column_width(table, (0.07, 0.45, 0.13, 0.35))
    
def add_list_methods(doc, name, class_info):
    
    doc.add_paragraph()
    doc.add_paragraph(name, style=doc.styles['Caption Table'])

    table = doc.add_table(rows=1, cols=4, style='Table H.Header wi.No')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # row 1
    table.rows[0].cells[0].text = 'No'
    table.rows[0].cells[1].text = 'Name'
    table.rows[0].cells[2].text = 'Type'
    table.rows[0].cells[3].text = 'Description'
    
    empty = 1
    
    #public method
    count = 0
    methods = class_info.PublicMethods
    for i in range(len(methods)):
        count += 1
        mt = methods[i]
        row = table.add_row()
        row.cells[0].text = to_asscii('(%i)' % (i+1))
        row.cells[1].text = to_asscii(mt.Name)
        row.cells[2].text = "Public"
        
        try:
            row.cells[3].text = to_asscii(mt.Description.lstrip())
        except:
            row.cells[3].text = 'None'
        empty = 0
        
    #protected method

    methods = class_info.ProtectedMethods
    for i in range(len(methods)):
        count += 1
        mt = methods[i]
        row = table.add_row()
        row.cells[0].text = to_asscii('(%i)' % (count))
        row.cells[1].text = to_asscii(mt.Name)
        row.cells[2].text = "Protected"
        
        try:
            row.cells[3].text = to_asscii(mt.Description.lstrip())
        except:
            row.cells[3].text = 'None'
        empty = 0
        
    #private method
    methods = class_info.PrivateMethods
    for i in range(len(methods)):
        count+= 1
        mt = methods[i]
        row = table.add_row()
        row.cells[0].text = to_asscii('(%i)' % (count))
        row.cells[1].text = to_asscii(mt.Name)
        row.cells[2].text = "Private"
        
        try:
            row.cells[3].text = to_asscii(mt.Description.lstrip())
        except:
            row.cells[3].text = 'None'
        empty = 0
        
    if empty == 1:
        table.add_row()
        table.rows[1].cells[0].text = '-'
        table.rows[1].cells[1].text = '-'
        table.rows[1].cells[2].text = '-'
        table.rows[1].cells[3].text = '-'

    #format_column_width(table, (0.1, 0.4, 0.5))
    table.autofit = False
    format_column_width(table, (0.07, 0.45, 0.13, 0.35))

def add_class_begin_document(document, class_info):
    table = document.add_table(rows=2, cols=2, style='Table V.Header')

    table.rows[0].cells[0].text = 'Location File Name'
    table.rows[0].cells[1].text = to_asscii(class_info.FileName)
    table.rows[1].cells[0].text = 'Description'
    try:
        table.rows[1].cells[1].text = to_asscii(class_info.Description.lstrip())
    except:
        table.rows[1].cells[1].text = 'None'

    format_column_width(table, (0.3, 0.7))

    document.add_paragraph()

class_tsdd = {}
class_field_tsdd = {}
def add_class_document(document, class_info):
    global class_count, field_method_count
    global method_count, tsdd, field_tsdd, class_tsdd, class_field_tsdd
    
    class_tsdd = {}
    class_field_tsdd = {}
    field_method_count = 0
    class_count += 1
    document.add_heading(class_info.Name, template_str.heading['classnamespace'])

    # Tables
    add_class_begin_document(document, class_info)

    fields = (
        ('Public Fields', class_info.PublicFields),
        ('Protected Fields', class_info.ProtectedFields),
        ('Private Fields', class_info.PrivateFields),
    )
    methods = (
        ('Public Methods', class_info.PublicMethods),
        ('Protected Methods', class_info.ProtectedMethods),
        ('Private Methods', class_info.PrivateMethods),
    )
    # for (key, value) in fields: add_list_methods(document, key, value)
    # for (key, value) in methods: add_list_methods(document, key, value)
    
    add_list_fields(document, "Field List", class_info)
    add_list_methods(document, "Method List", class_info)
    
    # Fields
    document.add_heading('Fields', template_str.heading['fieldstopic'])
    
    for (key, value) in fields:
        for mt in value:
            add_filed_document(document, mt)
    
    # Methods
    document.add_heading('Methods', template_str.heading['methodstopic'])
    
    for (key, value) in methods:
        for mt in sorted(value, key=lambda m:m.Name):
            add_method_document(document, mt)
    
    class_name = class_info.Name.split('::')[-1]
    tsdd[class_name] = class_tsdd
    field_tsdd[class_name] = class_field_tsdd
    
def print_help():
    print 'Usage python doc_gen.py xml_dir'
            
def adjust_picture_size(filepath):
    info = os.popen("file %s" % filepath).read()
    (x,y) = (float(z.strip()) for z in info.split(',')[1].split('x'))
    
    if float(x/72.0) > 6.0:
        return Inches(6.0)
    else:
        return None
            
def main():
    global REF_DB
    if (len(sys.argv)) > 1:
        src_path = sys.argv[1]
    else:
        print_help()
        exit()
    print 'Run with src dir: ', src_path

    #get configuration from json
    try:
        with open("configuration.json", 'r') as fjson:
            REF_DB = json.load(fjson)
    except:
        pass
    
    # Get py dir
    template_file_path = os.path.dirname(os.path.abspath(__file__)) + '/' + template_str.template_filepath
    document = Document(template_file_path)

    # temp = document.styles
    # for x in temp.element.style_lst:
    #     print x.name_val

    #document.add_heading('Class', template_str.heading['classname'])
    # Description
    #document.add_paragraph(template_str.diagram_comment)
    
    #get class infor with file name
    class_names = {}
    for root, dirs, files in os.walk(src_path):
        for file in files:
            if file.startswith('class') and 'map_name' not in file:
                filepath = (os.path.join(root, file))
                print ""
                print "[Process]", ntpath.basename(filepath)
                class_info = xmlparser.parse(filepath)
                if class_info is not None:
                    class_names[file.split('.')[0]] = class_info.Name
    
    # add charts
    for root, dirs, files in os.walk("html"):
        for file in files:
            if file.startswith('class') and 'map_name' not in file and file.endswith("inherit__graph.png"):
                filepath = (os.path.join(root, file)) 
                print ""
                print "[Process]", ntpath.basename(filepath)
                document.add_picture(filepath, width=adjust_picture_size(filepath))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph(class_names[file.split('__inherit__graph.png')[0]], style=document.styles['Class Caption'])
                del class_names[file.split('__inherit__graph.png')[0]]
                
        
    #find_class_image.find_image(document, class_names)
    for key in class_names:
        gen_graph = os.path.dirname(os.path.abspath(__file__)) + "/gen_graph.sh " 
        graph_name = class_names[key].strip().replace('::','.').strip()
        
        lg = len(graph_name)/2
        i = lg
        for i in range(lg,len(graph_name)):
            if graph_name[i] == '.':
                break
        lg = i+1       
        gn = "".join(graph_name[0:lg]) + "####" + "".join(graph_name[lg:])
        
        
        
        cmd = "bash %s %s %s" % (gen_graph, gn, graph_name)
        filepath = os.popen(cmd).read().strip()
        
        print filepath
        document.add_picture(filepath, width=adjust_picture_size(filepath))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(class_names[key], style=document.styles['Class Caption'])
    
    document.add_paragraph("")
    
    for root, dirs, files in os.walk(src_path):
        for file in files:
            if file.startswith('class') and 'map_name' not in file:
                filepath = (os.path.join(root, file))
                print ""
                print "[Process]", ntpath.basename(filepath)
                class_info = xmlparser.parse(filepath)
                if class_info is not None:
                    add_class_document(document, class_info)

    document.save(template_str.output_filepath)
    
    #Save gdd link to json
    with open ("gdd.json", 'w') as gddFile:
        json.dump(tsdd,  gddFile, indent=2)
    #Save field gdd link to json
    with open ("field_gdd.json", 'w') as gddFile:
        json.dump(field_tsdd,  gddFile, indent=2)
    print 'End'

def open_output_file():
    if os.name == 'posix':
        subprocess.call('cygstart ' + template_str.output_filepath, shell=True)
    else:
        subprocess.call('start ' + template_str.output_filepath, shell=True)

if __name__ == '__main__':
    main()
    open_output_file()
    print 'Have a nice day !'

