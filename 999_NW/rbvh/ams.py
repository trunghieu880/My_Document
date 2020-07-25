import __init__
import logging
import os
from pathlib import Path
from unicodedata import normalize
import lxml.etree
import lxml.html
import parse, utils
import datetime
import const as CONST
import subprocess
import re
from docx.api import Document
import win32timezone
import win32com.client
from win32com.client import Dispatch
import json, copy
import time
import random

logger = logging.getLogger(__name__)

class Base(object):
    def __init__(self, path):
        self.path = Path(path)

class FileTPA(Base):
    # Class FileTPA
    def __init__(self, path):
        super().__init__(path)
        self.doc = lxml.etree.parse(str(path))

    # Function get_tag: get the node of XML file with specific tag
    def get_tag(self, tag, index=0):
        '''Get normalized text of tag base on index of tag'''
        node = [e for e in self.doc.iterfind('.//{0}'.format(tag))][index]
        return node

    # Function update_tag: update the node with the specific tag and value
    def update_tag(self, tag, value):
        self.get_tag(tag).text = value
        path = re.sub("file:/", "", self.doc.docinfo.URL)
        with open(path, 'wb') as f:
            self.doc.write(f)

    # Function update_tpa: update tpa file with the input data
    def update_tpa(self, data):
        flag = False
        try:
            lst_header = ["UnitUnderTest", "ExecutionDate", "NTUserID", "FileName", "Verdict"]

            for h in lst_header:
                self.update_tag(h, data[h])

            # collect Percentage Coverage
            lst_Percentage = [e for e in self.doc.iterfind('.//{0}'.format("Percentage"))]

            lst_Percentage[0].text = data['C0']
            lst_Percentage[1].text = data['C1']
            if utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM":
                lst_Percentage[2].text = data['MCDCU']

            path = re.sub("file:/", "", self.doc.docinfo.URL)
            with open(path, 'wb') as f:
                self.doc.write(f)
            flag = True
        except Exception as e:
            flag = False
            logger.exception(e)
        finally:
            return flag

    # Function get_data: get the information in the Summary HTML file : "UnitUnderTest", "NTUserID", "FileName", "Verdict", "MetricName", "Percentage"
    def get_data(self):
        d = dict()
        try:
            lst_header = ["ExecutionDate", "UnitUnderTest", "NTUserID", "FileName", "Verdict", "MetricName", "Percentage"]
            ut = self.get_tag("UnitUnderTest").text
            user = self.get_tag("NTUserID").text
            func = self.get_tag("FileName").text
            result = self.get_tag("Verdict").text
            date = self.get_tag("ExecutionDate").text
            # collect title C0, C1, MC/DC
            lst_MetricName = [e.text for e in self.doc.iterfind('.//{0}'.format("MetricName"))]
            # collect Percentage Coverage
            lst_Percentage = [e.text for e in self.doc.iterfind('.//{0}'.format("Percentage"))]

            lst_data = list()
            lst_data.append("NTUserID")
            lst_data.append(user)
            lst_data.append("FileName")
            lst_data.append(func)
            lst_data.append("TestResult")
            lst_data.append(result)
            lst_data.append("ExecutionDate")
            lst_data.append(date)

            for index, key in enumerate(lst_MetricName):
                lst_data.append(key)
                lst_data.append(lst_Percentage[index])

            # convert list to json
            d[ut] = {lst_data[i]: lst_data[i + 1] for i in range(0, len(lst_data), 2)}
        except Exception as e:
            d = {}
            logger.exception(e)
        finally:
            return d

class FileTestReportXML(Base):
    # Class FileTestReportXML
    def __init__(self, path):
        super().__init__(path)
        self.doc = lxml.etree.parse(str(path))

    # Function get_tag: get the node of XML file with specific tag
    def get_tag(self, tag, index=0):
        '''Get normalized text of tag base on index of tag'''
        node = [e for e in self.doc.iterfind('.//{0}'.format(tag))][index]
        return node

    # Function get_data: get the information in the Summary HTML file : Verdict, C0, C1, MCDC
    def get_data(self):
        data = dict()
        try:
            lst_header = ["status", "statement", "decision", "booleanOperandEffectivenessMasking", "booleanOperandEffectivenessUnique", "testScriptName"]
            node_summary = self.get_tag("summary")
            status = {'Verdict': node_summary.attrib['status']}
            node_coverageInfo = self.get_tag("coverageInfo")[0]

            score = {'C0': item.text for item in node_coverageInfo if item.tag == "statement"}
            score = {**score, **{'C1': item.text for item in node_coverageInfo if item.tag == "decision"}}
            score = {**score, **{'MCDCM': item.text for item in node_coverageInfo if item.tag == "booleanOperandEffectivenessMasking"}}
            score = {**score, **{'MCDCU': item.text for item in node_coverageInfo if item.tag == "booleanOperandEffectivenessUnique"}}

            testscriptname = {'testScriptName': item.text for item in self.get_tag("info") if item.tag == "testScriptName"}

            data = {**status, **score, **testscriptname}
        except Exception as e:
            data = {}
            logger.exception(e)
        finally:
            return data

class FileTestReportHTML(Base):
    # Class FileTestReportHTML
    def __init__(self, path):
        super().__init__(path)
        self.doc = lxml.html.parse(str(path))

    # Function get_tag: get the node of html file with specific tag
    def get_tag(self, tag, index=0):
        '''Get normalized text of tag base on index of tag'''
        node = [e for e in self.doc.iterfind('.//{0}'.format(tag))][index]
        return node

    # T.B.D
    def get_data(self):
        pass

class FileTestSummaryHTML(Base):
    # Class FileTestSummaryHTML
    def __init__(self, path):
        super().__init__(path)
        self.doc = lxml.html.parse(str(path))

    # Function get_tag: get the node of html file with specific tag
    def get_tag(self, tag, index=0):
        '''Get normalized text of tag base on index of tag'''
        node = [e for e in self.doc.iterfind('.//{0}'.format(tag))][index]
        return node

    # Function get_data: get the information in the Summary HTML file : Verdict, C0, C1, MCDC
    def get_data(self):
        data = dict()
        try:
            for e in self.doc.iterfind('.//{0}'.format("div")):
                if e.text == None:
                    continue
                if "Project:" in e.text or "Overall Result:" in e.text:
                    for item in e:
                        if "Project:" in e.text:
                            data = {**data, **{'Project': item.text}}
                        elif "Overall Result:" in e.text:
                            data = {**data, **{'Verdict': item.text}}
                        else:
                            print("BUG FileTestSummaryHTML get_data")

            key = ""
            flag = False
            for e in self.doc.iterfind('.//{0}'.format("div")):
                if e.text == None:
                    continue
                if e.text == "Statement (S)" or e.text == "Decision (D)" or e.text == "MC/DC - masking (M)" or e.text == "MC/DC - unique cause (U)":
                    flag = True
                    if e.text == "Statement (S)":
                        key = "C0"
                    elif e.text == "Decision (D)":
                        key = 'C1'
                    elif e.text == "MC/DC - masking (M)":
                        key = 'MCDCM'
                    elif e.text == "MC/DC - unique cause (U)":
                        key = 'MCDCU'
                    else:
                        print("BUG")
                    next
                else:
                    if "%" in e.text and flag == True:
                        flag = False
                        val = e.text.replace("%", "")
                        data = {**data, key : val}
                        key = ""
                        val = ""
                        next
                    else:
                        continue

            flag = False
            for e in self.doc.iterfind('.//{0}'.format("div")):
                if e.text == None:
                    continue
                if e.text == "Summary generated":
                    flag = True
                    key = "date"
                    next
                else:
                    if flag == True:
                        flag = False
                        val = e.text
                        data = {**data, key : val}
                        break

        except Exception as e:
            data = {}
            logger.exception(e)
        finally:
            return data

def reformat_string(value):
    temp = re.sub("[\n\t\r\x07\xa0]", " ", value.strip()).strip()
    temp = re.sub("\s+", " ", temp).strip()
    return temp

class FileWalkThroughDoc(Base):
    def __init__(self, path):
        super().__init__(path)
        self.doc = str(path)

    # Function update_tpa: update tpa file with the input data
    def update(self, data, opt="PSW"):
        flag = False
        try:
            if opt == "PSW":
                document = Document(self.doc)
                table_infor = document.tables[1]
                table_infor.cell(0,1).text = data['date']
                table_infor.cell(0,3).text = data['project']
                table_infor.cell(0,5).text = data['review initiator']
                table_infor.cell(1,1).text = str(data['effort'])
                table_infor.cell(1,3).text = data['baseline']
                table_infor.cell(1,5).text = data['review partner']
                table_attach = document.tables[2]
                table_attach.cell(1, 2).text = data['path_testscript']
                table_attach.cell(3, 2).text = data['path_test_summary']
                table_attach.cell(3, 1).text = data['ScoreC0C1']
                document.save(self.doc)
            elif opt == "ASW":
                pass
                # word = win32com.client.DispatchEx('Word.Application')
                # word.Visible = 0
                # word.DisplayAlerts = 0

                # doc = word.Documents.Open(self.doc)

                # table_infor = doc.Tables(2)
                # table_attach = doc.Tables(3)
                # table_finding = doc.Tables(4)
                # table_check_list = doc.Tables(8)

                # finding = reformat_string(table_finding.Cell(Row=2, Column=2).Range.Text)
                # impact = reformat_string(table_finding.Cell(Row=2, Column=4).Range.Text)
                # confirm_UT26 = reformat_string(table_check_list.Cell(Row=12, Column=5).Range.Text)

                # temp = reformat_string(table_attach.Cell(Row=7, Column=2).Range.Text)
                # [score_c0, score_c1, score_mcdc] = re.sub("^.*C0: ([0-9]+)%.*C1: ([0-9]+)%.*MCDC: ([0-9]+)%", r'\1 \2 \3', temp).split(" ")

                # dict_walkthrough = {
                #     'date': reformat_string(table_infor.Cell(Row=1, Column=2).Range.Text),
                #     'project': reformat_string(table_infor.Cell(Row=1, Column=4).Range.Text),
                #     'review initiator': reformat_string(table_infor.Cell(Row=1, Column=6).Range.Text),
                #     'effort': reformat_string(table_infor.Cell(Row=2, Column=2).Range.Text),
                #     'baseline': reformat_string(table_infor.Cell(Row=2, Column=4).Range.Text),
                #     'review partner' : reformat_string(table_infor.Cell(Row=2, Column=6).Range.Text),
                #     'C0': score_c0,
                #     'C1': score_c1,
                #     'MCDC': score_mcdc,
                #     'tbl_finding': {
                #         "finding": finding,
                #         "impact": impact,
                #         "confirm_UT26": confirm_UT26
                #     }
                # }

                # doc.Save()
                # doc.Close()
                # word.Quit()
                # word = None
            else:
                logger.error("None Bug Type")
            
            flag = True
        except Exception as e:
            flag = False
            logger.exception(e)
        finally:
            return flag

    # Function get_data: to get the array data with specfic key, value of the nested json
    def get_data(self, opt="PSW"):
        dict_walkthrough = dict()
        try:
            if opt == "PSW":
                document = Document(self.doc)
                table_infor = document.tables[1]
                table_attach = document.tables[2]
                table_finding = document.tables[3]
                table_check_list = document.tables[6]

                temp = re.sub("[\n\t]", " ", table_attach.cell(3, 1).text).strip()
                [score_c0, score_c1] = re.sub("^.*C0: ([0-9]+).*C1: ([0-9]+)", r'\1 \2', temp).replace("%", "").split(" ")

                finding = table_finding.cell(1, 1).text
                impact = table_finding.cell(1, 3).text
                confirm_UT9 = table_check_list.cell(12, 3).text

                dict_walkthrough = {
                    'date': table_infor.cell(0,1).text,
                    'project': table_infor.cell(0,3).text,
                    'review initiator': table_infor.cell(0,5).text,
                    'effort': table_infor.cell(1,1).text,
                    'baseline': table_infor.cell(1,3).text,
                    'review partner' : table_infor.cell(1,5).text,
                    'path_testscript': table_attach.cell(1, 2).text,
                    'path_test_summary': table_attach.cell(3, 2).text,
                    'C0': score_c0,
                    'C1': score_c1,
                    'tbl_finding': {
                        "finding": finding,
                        "impact": impact,
                        "confirm_UT9": confirm_UT9
                    }
                }
                
            elif opt == "ASW":
                if(re.search(".docx$", self.doc)):
                    doc = Document(self.doc)
                    table_infor = doc.tables[1]
                    table_attach = doc.tables[2]
                    table_finding = doc.tables[3]
                    table_check_list = doc.tables[7]
                    
                    finding = reformat_string(table_finding.cell(1, 1).text)
                    impact = reformat_string(table_finding.cell(1, 3).text)
                    confirm_UT26 = reformat_string(table_check_list.cell(11, 4).text)

                    temp = reformat_string(table_finding.cell(2, 1).text)
                    score_c0 = score_c1 = score_mcdc = ""
                    if re.search("^.*C0: ([0-9]+).*C1: ([0-9]+).*MCDC: (.*)", temp):
                        temp = re.sub("^.*C0: ([0-9]+).*C1: ([0-9]+).*MCDC: (.*)", r'\1 \2 \3', temp).replace("%", "").split(" ")
                        score_c0 = temp[0]
                        score_c1 = temp[1]
                        score_mcdc = temp[2]
                    else:
                        score_c0 = score_c1 = score_mcdc = "Unknown"
                        logger.error(">> Failed: Item FileWalkThroughDoc has wrong format at Table Review Item: Coverage Report: \\nC0: <>%\\nC1: <>%\\nMCDC: <>%")

                    temp = reformat_string(table_infor.cell(0, 3).text)
                    if re.search("^(.*) \<V(.*)\>", temp):
                        temp = re.sub("^(.*) \<V(.*)\>", r'\1 \2', temp).split(" ")
                        project_name = temp[0]
                        item_revision = temp[1]
                    else:
                        project_name = "Unknown"
                        item_revision = "Unknown"
                        logger.error(">> Failed: Item FileWalkThroughDoc has wrong format project: [Class] <V[Version]>")

                    dict_walkthrough = {
                        'date': reformat_string(table_infor.cell(0, 1).text),
                        'project': project_name,
                        'ItemRevision': item_revision.replace("_", "."),
                        'review initiator': reformat_string(table_infor.cell(0, 5).text),
                        'effort': reformat_string(table_infor.cell(1, 1).text),
                        'baseline': reformat_string(table_infor.cell(1, 3).text),
                        'review partner' : reformat_string(table_infor.cell(1, 5).text),
                        'C0': score_c0,
                        'C1': score_c1,
                        'MCDC': score_mcdc,
                        'tbl_finding': {
                            "finding": finding,
                            "impact": impact,
                            "confirm_UT26": confirm_UT26
                        }
                    }
                else:
                    word = win32com.client.DispatchEx('Word.Application')
                    word.Visible = 0
                    word.DisplayAlerts = 0

                    doc = word.Documents.Open(self.doc)
                    table_infor = doc.Tables(2)
                    table_attach = doc.Tables(3)
                    table_finding = doc.Tables(4)
                    table_check_list = doc.Tables(8)

                    
                    finding = reformat_string(table_finding.Cell(Row=2, Column=2).Range.Text)
                    impact = reformat_string(table_finding.Cell(Row=2, Column=4).Range.Text)
                    confirm_UT26 = reformat_string(table_check_list.Cell(Row=12, Column=5).Range.Text)

                    temp = reformat_string(table_finding.Cell(Row=3, Column=2).Range.Text)
                    score_c0 = score_c1 = score_mcdc = ""
                    if re.search("^.*C0: ([0-9]+).*C1: ([0-9]+).*MCDC: (.*)", temp):
                        temp = re.sub("^.*C0: ([0-9]+).*C1: ([0-9]+).*MCDC: (.*)", r'\1 \2 \3', temp).replace("%", "").split(" ")
                        score_c0 = temp[0]
                        score_c1 = temp[1]
                        score_mcdc = temp[2]
                    else:
                        score_c0 = score_c1 = score_mcdc = "Unknown"
                        logger.error(">> Failed: Item FileWalkThroughDoc has wrong format at Table Review Item: Coverage Report: \\nC0: <>%\\nC1: <>%\\nMCDC: <>%")

                    temp = reformat_string(table_infor.Cell(Row=1, Column=4).Range.Text)
                    if re.search("^(.*) \<V(.*)\>", temp):
                        temp = re.sub("^(.*) \<V(.*)\>", r'\1 \2', temp).split(" ")
                        project_name = temp[0]
                        item_revision = temp[1]
                    else:
                        project_name = "Unknown"
                        item_revision = "Unknown"
                        logger.error(">> Failed: Item FileWalkThroughDoc has wrong format project: [Class] <V[Version]>")

                    dict_walkthrough = {
                        'date': reformat_string(table_infor.Cell(Row=1, Column=2).Range.Text),
                        'project': project_name,
                        'ItemRevision': item_revision.replace("_", "."),
                        'review initiator': reformat_string(table_infor.Cell(Row=1, Column=6).Range.Text),
                        'effort': reformat_string(table_infor.Cell(Row=2, Column=2).Range.Text),
                        'baseline': reformat_string(table_infor.Cell(Row=2, Column=4).Range.Text),
                        'review partner' : reformat_string(table_infor.Cell(Row=2, Column=6).Range.Text),
                        'C0': score_c0,
                        'C1': score_c1,
                        'MCDC': score_mcdc,
                        'tbl_finding': {
                            "finding": finding,
                            "impact": impact,
                            "confirm_UT26": confirm_UT26
                        }
                    }
            else:
                dict_walkthrough = {}
                logger.error("None Bug Type")
        except Exception as e:
            dict_walkthrough = {}
            logger.exception(e)
        finally:
            # if opt == "ASW":
            #     doc.Close()
            #     word.Quit()
            #     word = None
            return dict_walkthrough

class FileCoverageReasonXLS(Base):
    def __init__(self, path):
        super().__init__(path)
        self.doc = str(path)

    # Function update_tpa: update tpa file with the input data
    def update(self, data):
        flag = False
        try:
            excel = win32com.client.Dispatch('Excel.Application')
            wb = excel.Workbooks.Open(self.doc)
            excel.Visible = False
            excel.DisplayAlerts = False
            wb.DoNotPromptForConvert = True
            wb.CheckCompatibility = False

            score_c0 = (value(formatNumber(float(value(data.get("C0"))) * 100)) if (value(data.get("C0")) != "-" and data.get("C0") != None) else "NA")
            score_c1 = (value(formatNumber(float(value(data.get("C1"))) * 100)) if (value(data.get("C1")) != "-" and data.get("C1") != None) else "NA")
            score_mcdc = (value(formatNumber(float(value(data.get("MCDC"))) * 100)) if (value(data.get("MCDC")) != "-" and data.get("MCDC") != None) else "NA")

            data_tpa = {
                "UnitUnderTest": data.get("ItemName"),
                "NTUserID": str(convert_name(key=data.get("Tester"), opt="id")),
                "ExecutionDate" : datetime.datetime.now().strftime("%Y-%m-%d"),
                "C0": score_c0,
                "C1": score_c1,
                "MCDCU": score_mcdc
            }

            writeData = wb.Worksheets(1)
            # Write data here
            infor_CoverageReasonXLS = utils.load(CONST.SETTING).get("CoverageReasonXLS")

            writeData.Range(infor_CoverageReasonXLS.get("Tester")).Value = data_tpa.get("NTUserID")
            writeData.Range(infor_CoverageReasonXLS.get("Date")).Value = data_tpa.get("ExecutionDate")
            writeData.Range(infor_CoverageReasonXLS.get("Item_Name")).Value = data_tpa.get("UnitUnderTest")
            writeData.Range(infor_CoverageReasonXLS.get("C0")).Value = data_tpa.get("C0")
            writeData.Range(infor_CoverageReasonXLS.get("C1")).Value = data_tpa.get("C1")
            writeData.Range(infor_CoverageReasonXLS.get("MCDC")).Value = data_tpa.get("MCDCU")

            flag = True
        except Exception as e:
            flag = False
            logger.exception(e)
        finally:
            wb.Save()
            wb.Close()
            excel.Quit()
            excel = None
            return flag

    # Function get_data: to get the array data with specfic key, value of the nested json
    def get_data(self):
        data = dict()
        try:
            excel = win32com.client.Dispatch('Excel.Application')
            wb = excel.Workbooks.Open(self.doc)

            excel.Visible = False
            excel.DisplayAlerts = False
            wb.DoNotPromptForConvert = True
            wb.CheckCompatibility = False

            readData = wb.Worksheets(1)
            allData = readData.UsedRange

            infor_CoverageReasonXLS = utils.load(CONST.SETTING).get("CoverageReasonXLS")

            # data = {
            #     "Tester": value(allData.Cells(1, 2).value),
            #     "Date": value(allData.Cells(2, 2).value),
            #     "Item_Name": value(allData.Cells(3, 2).value),
            #     "C0": value(formatNumber(float(allData.Cells(9, 2).value))),
            #     "C1": value(formatNumber(float(allData.Cells(10, 2).value))),
            #     "MCDC": value(formatNumber(float(allData.Cells(11, 2).value)))
            # }

            data = {
                "Tester": value(allData.Cells(1, 2)),
                "Date": value(allData.Cells(2, 2)),
                "Item_Name": value(allData.Cells(3, 2)),
                "C0": value(formatNumber(float(value(allData.Cells(9, 2))))),
                "C1": value(formatNumber(float(value(allData.Cells(10, 2))))),
                "MCDC": value(formatNumber(float(value(allData.Cells(11, 2)))))
            }

        except Exception as e:
            data = {}
            logger.exception(e)
        finally:
            wb.Save()
            wb.Close()
            excel.Quit()
            excel = None
            return data

class FileSummaryXLSX(Base):
    # Class FileSummaryXLSX
    def __init__(self, path):
        super().__init__(path)
        self.doc = str(path)

    # Function parse2json: convert XLSX to json data
    def parse2json(self, sheetname="", begin=59, end=59):
        return dict(parse.parse_summary_json(self.doc, sheetname=sheetname, begin=begin, end=end))

    # Function get_data: to get the array data with specfic key, value of the nested json
    def get_data(self, data, key, value):
        result = dict()
        try:
            item = -1
            result = dict()
            for item in data.keys():
                if(data[item].get(key) == value):
                    result[item] = data[item]
        except Exception as e:
            result = {}
            logger.exception(e)
        finally:
            return result

class FileRTRTCov(Base):
    # Class FileRTRTCov
    def __init__(self, path):
        super().__init__(path)
        self.doc = str(path)

    # Function get_data: to get the array data with specfic key, value of the nested json
    def get_data(self):
        try:
            flag_count = 0

            data = dict()
            with open(self.doc, encoding='shift-jis', errors='ignore') as fp:
                for line in fp.readlines()[:100]:
                    line = line.strip()
                    if flag_count > 1:
                        break

                    if 'Conclusion :' in line:
                        flag_count = flag_count + 1
                        next
                    if (flag_count == 1):
                        if 'Statement blocks' in line or 'Decisions' in line or 'Modified conditions' in line:
                            temp = re.sub("\s+\.+\s+", ":", line)
                            temp = re.sub("%\s\(.*\)$", "", temp)
                            key = temp.split(":")[0]
                            val = convert_score_percentage(temp.split(":")[1], opt='nomul')
                            if key == "Statement blocks":
                                key = "C0"
                            elif key == "Decisions":
                                key = 'C1'
                            elif key == "Modified conditions":
                                key = 'MCDC'
                                rst = True

                            data = {**data, key : val}

        except Exception as e:
            data = {}
            logger.exception(e)
        finally:
            return data

class FileTestDesignXLSX(Base):
    # Class FileTestDesignXLSX
    def __init__(self, path):
        super().__init__(path)
        self.doc = str(path)

    # Function parse2json: convert XLSX to json data
    def parse2json(self, sheetname, begin=59, end=59):
        ROW_ENUM = 15 - 1
        ROW_TOLERANCE = 19 - 1
        ROW_HEADER_INPUT = ROW_TOLERANCE + 4
        ROW_NAME_VARIABLE = ROW_HEADER_INPUT + 1

        if utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM":
            ROW_ENUM = ROW_ENUM - 1
            ROW_TOLERANCE = ROW_TOLERANCE - 1
            ROW_HEADER_INPUT = ROW_TOLERANCE + 4
            ROW_NAME_VARIABLE = ROW_HEADER_INPUT + 1

        def check_flag(list_data, pat):
            result = False
            for x in list_data:
                if x == pat:
                    result = True

            return result

        def initial_value(header):
            result = dict()
            for __header__ in header:
                result[__header__] = {"begin": -1, "end": -1}

            return result

        def find_cordinate(result, data, header, new_header=[], isEnum=False, row=0):
            for index, __header__ in enumerate(header):
                result[__header__]['begin'] = data[row].index(__header__)
                if (__header__ == header[-1]):
                    if isEnum == True:
                        list_temp = [x for x in data[row + 1] if (x is not None)]
                        result[__header__]['end'] = data[row + 1].index(list_temp[-1])
                    else:
                        result[__header__]['end'] = data[row].index(header[index])
                else:
                    result[__header__]['end'] = data[row].index(header[index + 1]) - 1

            for key in dict(result).keys():
                if isEnum == True:
                    if (key == "Enum"):
                        del result[key]
                else:
                    if not(key in new_header):
                        del result[key]

        def generate_data(data_header, data, row, header=[], isEnum=False):
            result = dict()
            if isEnum == True:
                for __header__ in dict(data_header).keys():
                    result[__header__] = list()
                    for j in range(data_header[__header__]['begin'], data_header[__header__]['end'] + 1):
                        result[__header__].append(data[row][j])
                
            else:
                template_obj = {
                    "Tolerance": -1,
                    "Type": None,
                    "Max": -1,
                    "Min": -1,
                    "Data": list()
                }
            
                max_testcase = -1
                for i in range(ROW_NAME_VARIABLE + 1, len(data)):
                    if(i == len(data) - 1):
                        max_testcase = i - (ROW_NAME_VARIABLE + 1) + 1
                        break
                    else:
                        if (data[i][0] is None or ("TC" in data[i][0]) == False):
                            max_testcase = i - (ROW_NAME_VARIABLE + 1)
                            break

                for __header__ in header:
                    result[__header__] = list()
                    flag_end_data = False
                    for j in range(data_header[__header__]['begin'], data_header[__header__]['end'] + 1):
                        name_input = data[ROW_NAME_VARIABLE][j]
                        if __header__ == "DESCRIPTIONS":
                            name_input = 'DESCRIPTIONS'
                        
                        data_js_variable = copy.deepcopy(template_obj)
                        
                        data_js_variable["Tolerance"] = data[ROW_TOLERANCE][j]
                        data_js_variable["Type"] = data[ROW_TOLERANCE + 1][j]
                        data_js_variable["Max"] = data[ROW_TOLERANCE + 2][j]
                        data_js_variable["Min"] = data[ROW_TOLERANCE + 3][j]
                        
                        for i in range(ROW_NAME_VARIABLE + 1, ROW_NAME_VARIABLE + 1 + max_testcase):
                            data_js_variable["Data"].append(data[i][j])

                        result[__header__].append({name_input: data_js_variable})
            return result
        
        full_data = parse.get_xlsx_raw(self.doc, sheet=sheetname, begin=begin, end=end)
        header_input = ['INPUTS', 'LOCAL VARIABLES AS INPUT', 'IMPORTED PARAMETERS', 'LOCAL PARAMETERS', 'LOCAL VARIABLES', 'OUTPUTS', 'DESCRIPTIONS']
        data_header = full_data[ROW_HEADER_INPUT]
        finding_header = [x for x in data_header if (x is not None and (not re.search("^#", value(x))))]
        
        for __header__ in list(header_input):
            if not (__header__ in finding_header):
                header_input.remove(__header__)

        ''' Count element '''
        d_input_header = dict()
        d_input_header = initial_value(header=finding_header)
        find_cordinate(result=d_input_header, data=full_data, header=finding_header, new_header=header_input, row=ROW_HEADER_INPUT)
        result = dict()
        result = generate_data(data_header=d_input_header, header=header_input, data=full_data, row=ROW_TOLERANCE)

        data_enum_title = full_data[ROW_ENUM]
        enum_header = [x for x in data_enum_title if ((x is not None) and (x != "Enum"))]
        d_enum_header = dict()
        d_enum_header = initial_value(header=enum_header)
        find_cordinate(result=d_enum_header, data=full_data, header=enum_header, isEnum=True, row=ROW_ENUM)
        result_enum = dict()
        result_enum = generate_data(data_header=d_enum_header, data=full_data, isEnum=True, row=ROW_ENUM + 1)
        
        final_result = dict()
        final_result = {
            "sheetname": sheetname,
            "Header": header_input,
            "Enum": result_enum,
            "TestDesign": result
        }

        return final_result

    # Function get_data: to get the array data with specfic key, value of the nested json
    def get_data(self):
        data = dict()
        BEGIN=1
        END=5000
        ROW_ENUM = 15 - 1
        ROW_TOLERANCE = 19 - 1
        ROW_HEADER_INPUT = ROW_TOLERANCE + 4
        ROW_NAME_VARIABLE = ROW_HEADER_INPUT + 1
        if utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM":
            ROW_ENUM = ROW_ENUM - 1
            ROW_TOLERANCE = ROW_TOLERANCE - 1
            ROW_HEADER_INPUT = ROW_TOLERANCE + 4
            ROW_NAME_VARIABLE = ROW_HEADER_INPUT + 1
        
        TM_CELL = "A{}".format(ROW_NAME_VARIABLE + 1)

        try:
            if utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM":
                item_name = re.sub("^TD_(\w.*)_(MT_\d.*)\.xls.*$", r'\1 \2', os.path.basename(self.doc)).split(" ")[0]
                item_revision = "NA"
            else:
                try:
                    temp =  re.sub("^TD_(\w.*)_v(\d.*)\.xls.*$", r'\1 \2', os.path.basename(self.doc)).split(" ")
                    item_name = temp[0]
                    item_revision = temp[1]
                except Exception as e:
                    item_name = "Uknown"
                    item_revision = "Unknown"
                    logger.error(">> Failed: Item FileTestDesignXLSX has wrong format name: {}".format(os.path.basename(self.doc)))
                    # logger.exception(e)


            ignore_sheet = ['Revision History', 'Summary', 'MCDC', 'Temporary', 'Guide', 'SavedVariables', 'MCDC_Backup', 'Testcases_Backup', 'Constants_Backup']
            lst_sheet = [x for x in parse.get_xlsx_sheets(self.doc) if (x.strip() not in ignore_sheet and parse.get_xlsx_raw(self.doc, sheet=x, begin=BEGIN, end=END)[ROW_NAME_VARIABLE + 1][0] == "TC1")]
            
            sheet_history = [x for x in parse.get_xlsx_sheets(self.doc) if (x.strip() == "Revision History")][0]

            data_test_design = list()

            for sheet in lst_sheet:
                data_test_design.append(self.parse2json(sheetname=sheet, begin=BEGIN, end=END))

            data = {
                "ItemName": item_name,
                "TM": [parse.get_xlsx_cells(xlsx=self.doc, sheet=x, list_cell=[TM_CELL]).get(TM_CELL) for x in lst_sheet],
                "ItemRevision": item_revision.replace("_", "."),
                "Tester": parse.get_xlsx_cells(xlsx=self.doc, sheet=sheet_history, list_cell=['C17']).get('C17'),
                "Date": parse.get_xlsx_cells(xlsx=self.doc, sheet=sheet_history, list_cell=['D17']).get('D17'),
                "DataTestDesign": data_test_design
            }

        except Exception as e:
            data = {}
            logger.exception(e)
        finally:
            return data

class FileATTReportXML(Base):
    # Class FileTestReportXML
    def __init__(self, path):
        super().__init__(path)
        self.doc = lxml.etree.parse(str(path))

    # Function get_tag: get the node of XML file with specific tag
    def get_tag(self, tag, index=0):
        '''Get normalized text of tag base on index of tag'''
        node = [e for e in self.doc.iterfind('.//{0}'.format(tag))][index]
        return node

    # Function get_data: get the information in the Summary HTML file : Verdict, C0, C1, MCDC
    def get_data(self):
        data = dict()
        try:
            lst_header = ["ClassName", "ClassVersion", "CompleteVerdict", "TestModuleName"]
            for index, key in enumerate(lst_header):
                if key == 'TestModuleName':
                    temp_lst_test_module = [e.text for e in self.doc.iterfind('.//{0}'.format("TestModuleName"))]
                    data = {**data, **{key: temp_lst_test_module}}
                else:
                    data = {**data, **{key: self.get_tag(key).text}}
            
        except Exception as e:
            data = {}
            logger.exception(e)
        finally:
            return data

# Check the directory of function is exist or not
def check_exist(dir_input, function):
    return Path(dir_input).joinpath(function).exists()

# Convert Tester name to Real Name or USERID
def convert_name(key, opt="name"):
    logger.debug("Convert name")
    try:
        users = utils.load(CONST.SETTING).get("users")

        if opt == "name" or opt == "id":
            return str(users[key].get(opt))
        else:
            logger.error("Bug convert name")
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Check path : have src extension or not
def is_have_src(path):
    return ("\\src" in path or "\\SRC" in path)

# Trim Component Path to find the correct path
def trim_src(path):
    result = re.sub("^.*\\\\rb\\\\", "rb\\\\", path)
    result = re.sub("^.*\\\\rba\\\\", "rba\\\\", result)
    result = re.sub("\\\\src\\\\.*$", "", result)
    result = re.sub("\\\\SRC\\\\.*$", "", result)
    result = re.sub("\\\\src$", "", result)
    result = re.sub("\\\\SRC$", "", result)
    return result

def value(cell):
    if str(cell).isdigit():
        return str(cell)
    else:
        return str(cell)

def convert_score_percentage(num, opt=""):
    try:
        if (utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM"):
            if num == "NA" or num == "none":
                return "NA"

        if opt == 'nomul':
            return value(formatNumber(float(value(num))))
        else:
            return value(formatNumber(round(float(value(formatNumber(float(value(num)) * 100))), 1)))
    except ValueError:
        return "Unknown"

def check_score(score_test_summary, score_exel, opt=""):
    return ((score_test_summary == convert_score_percentage(score_exel) if (value(score_exel) != "-" and score_exel != None) else "NA"))

def formatNumber(num):
    return int(num) if num % 1 == 0 else num

# Check the number attachment of OPL
def check_OPL_Walkthrough(file):
    num_OPL = -1
    try:
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = 0
        word.DisplayAlerts = 0

        doc = word.Documents.Open(file.as_posix())

        num_OPL =  doc.InlineShapes.Count - 1
    except Exception as e:
        num_OPL = -1
        logger.error("Walkthrough can not open!!!. Please check your name carefully at link: {}".format(file))
        # logger.exception(e)
    finally:
        doc.Close()
        word.Quit()
        word = None
        return num_OPL

# Check def check_DataTestDesignXLSX(data):
def check_DataTestDesignXLSX(data_test_design):
    flag = True
    logger.debug("Function check_DataTestDesignXLSX")
    try:
        for __header__ in data_test_design['Header']:
            if __header__ == "LOCAL PARAMETERS" or __header__ == 'LOCAL VARIABLES' or __header__ == 'OUTPUTS':
                continue

            for data_var in data_test_design['TestDesign'][__header__]:
                for key, obj_level_2 in data_var.items():
                    list_data_var = list(set(obj_level_2["Data"]))
                    if not utils.isMissingValue(list_data_var):
                        list_data_var.sort(key=None, reverse=False)
                        if __header__ == 'DESCRIPTIONS':
                            logger.info(">> Passed: Your comment is missing: {}".format(key))
                        else:
                            logger.info(">> Passed: Your input is missing or wrong type value: {} <=> {}".format(key, list_data_var))
                    else:
                        flag = False
                        if __header__ == 'DESCRIPTIONS':
                            logger.error(">> Failed: Your comment is missing: {}".format(key))
                        else:
                            logger.error(">> Failed: Your input is missing or wrong type value: {} <=> {}".format(key, list_data_var))
                        continue
                    
                    if __header__ == 'DESCRIPTIONS':
                        continue

                    if obj_level_2['Type'] == 'log':
                        if not (False in list_data_var and True in list_data_var and len(list_data_var) == 2):
                            flag = False
                            logger.error("Header {}, Variable {}, Type {}: {} <=> {}".format(__header__, key, obj_level_2['Type'], "[False, True]", list_data_var))
                    elif obj_level_2['Type'] == 'cont' or obj_level_2['Type'] == 'sdisc' or obj_level_2['Type'] == 'udisc':
                        if __header__ == "INPUTS" and utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM":
                            if not (max(list_data_var) > obj_level_2['Max']):
                                flag = False
                                logger.error(">> Failed: Header {}, Variable {}, Type {}: Missing OUT RANGE MAX: {} <=> {}".format(__header__, key, obj_level_2['Type'], obj_level_2['Max'], max(list_data_var)))
                            else:
                                logger.info(">> Passed: Header {}, Variable {}, Type {}: Missing OUT RANGE MAX: {} <=> {}".format(__header__, key, obj_level_2['Type'], obj_level_2['Max'], max(list_data_var)))

                            if not (min(list_data_var) < obj_level_2['Min']):
                                flag = False
                                logger.error(">> Failed: Header {}, Variable {}, Type {}: Missing OUT RANGE MIN: {} <=> {}".format(__header__, key, obj_level_2['Type'], obj_level_2['Min'], min(list_data_var)))
                            else:
                                logger.info(">> Passed: Header {}, Variable {}, Type {}: Missing OUT RANGE MIN: {} <=> {}".format(__header__, key, obj_level_2['Type'], obj_level_2['Min'], min(list_data_var)))
                        else:

                            if __header__ == "IMPORTED PARAMETERS" and ("INF" in str(obj_level_2['Max']) or "inf" in str(obj_level_2['Max'])):
                                if not (len(list_data_var) >= 3):
                                    flag = False
                                    logger.error(">> Failed: Header {}, Variable {}, Type {}: Missing MAX/MIN/MID: {} <=> {}".format(__header__, key, obj_level_2['Type'], "+-{}".format(obj_level_2['Max']), list_data_var))
                                else:
                                    logger.info(">> Passed: Header {}, Variable {}, Type {}: Missing MAX/MIN/MID: {} <=> {}".format(__header__, key, obj_level_2['Type'], "+-{}".format(obj_level_2['Max']), list_data_var))
                                
                                continue
                            else:
                                if not (obj_level_2['Max'] in list_data_var):
                                    flag = False
                                    logger.error(">> Failed: Header {}, Variable {}, Type {}: Missing MAX: {} <=> {}".format(__header__, key, obj_level_2['Type'], obj_level_2['Max'], list_data_var))
                                else:
                                    logger.info(">> Passed: Header {}, Variable {}, Type {}: Missing MAX: {} <=> {}".format(__header__, key, obj_level_2['Type'], obj_level_2['Max'], list_data_var))
                                if not (obj_level_2['Min'] in list_data_var):
                                    flag = False
                                    logger.error(">> Failed: Header {}, Variable {}, Type {}: Missing MIN: {} <=> {}".format(__header__, key, obj_level_2['Type'], obj_level_2['Min'], list_data_var))
                                else:
                                    logger.info(">> Passed: Header {}, Variable {}, Type {}: Missing MIN: {} <=> {}".format(__header__, key, obj_level_2['Type'], obj_level_2['Min'], list_data_var))
                            
                        check_mid = [x for x in list_data_var if (x != obj_level_2['Max'] and x != obj_level_2['Min'])]
                        
                        tolerance = obj_level_2['Tolerance']
                        if __header__ == "IMPORTED PARAMETERS":
                            tolerance = 0

                        if obj_level_2['Max'] > obj_level_2['Min']:
                            if ((obj_level_2['Max'] - tolerance) != obj_level_2['Min']):
                                if not (len(check_mid) != 0 and utils.checkMidValue(lst=check_mid, min=obj_level_2['Min'], max=obj_level_2['Max'])):
                                    flag = False
                                    logger.error(">> Failed: Header {}, Variable {}, Type {}: Missing mid {}".format(__header__, key, obj_level_2['Type'], check_mid))
                                else:
                                    logger.info(">> Passed: Header {}, Variable {}, Type {}: Missing mid {}".format(__header__, key, obj_level_2['Type'], check_mid))
                            else:
                                if not (len(check_mid) == 0):
                                    flag = False
                                    logger.error(">> Failed: Header {}, Variable {}, Type {}: None Mid because MAX = MIN".format(__header__, key, obj_level_2['Type']))
                                else:
                                    logger.info(">> Passed: Header {}, Variable {}, Type {}: None Mid because MAX = MIN".format(__header__, key, obj_level_2['Type']))
                        elif obj_level_2['Max'] == obj_level_2['Min']:
                            if not (len(check_mid) == 0):
                                flag = False
                                logger.error(">> Failed: Header {}, Variable {}, Type {}: May be constant {}".format(__header__, key, obj_level_2['Type'], obj_level_2['Max']))
                            else:
                                logger.info(">> Passed: Header {}, Variable {}, Type {}: May be constant {}".format(__header__, key, obj_level_2['Type'], obj_level_2['Max']))
                        else:
                            flag = False
                            logger.error(">> Failed: Header {}, Variable {}, Type {}: BUG MAX NEVER LESS THAN MIN".format(__header__, key, obj_level_2['Type']))

                    elif obj_level_2['Type'] in data_test_design['Enum'].keys() or obj_level_2['Type'] == 'enum':
                        value_find_type_enum = obj_level_2['Type']
                        current_type = random.choice(list_data_var)
                        if obj_level_2['Type'] == 'enum' and utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM":
                            value_find_type_enum = utils.find_enum(pat=current_type, data=data_test_design['Enum']) 
                            if not(value_find_type_enum is not None):
                                flag = False
                                logger.error(">> Failed: Header {}, Variable {}, Type ENUM with value {}: Can not find your type".format(__header__, key, current_type))
                                continue
                            else:
                                logger.info(">> Passed: Header {}, Variable {}, Type ENUM with value {}: Can not find your type".format(__header__, key, current_type))

                        if not (data_test_design['Enum'][value_find_type_enum][-1] in list_data_var):
                            if data_test_design['Enum'][value_find_type_enum][-1] is None:
                                flag = False
                                logger.error(">> Failed: Header {}, Variable {}, Type ENUM {}: Please check your enum header that it hasn't any redundant table at here: {}".format(__header__, key, value_find_type_enum, list_data_var))
                            else:
                                flag = False
                                logger.error(">> Failed: Header {}, Variable {}, Type ENUM {}: Missing MAX {} <=> {}".format(__header__, key, value_find_type_enum, data_test_design['Enum'][value_find_type_enum][-1], list_data_var))
                        else:
                            logger.info(">> Passed: Header {}, Variable {}, Type ENUM {}: Missing MAX {} <=> {}".format(__header__, key, value_find_type_enum, data_test_design['Enum'][value_find_type_enum][-1], list_data_var))
                        if not (data_test_design['Enum'][value_find_type_enum][0] in list_data_var):
                            flag = False
                            logger.error(">> Failed: Header {}, Variable {}, Type ENUM {}: Missing MIN {} <=> {}".format(__header__, key, value_find_type_enum, data_test_design['Enum'][value_find_type_enum][0], list_data_var))
                        else:
                            logger.info(">> Passed: Header {}, Variable {}, Type ENUM {}: Missing MIN {} <=> {}".format(__header__, key, value_find_type_enum, data_test_design['Enum'][value_find_type_enum][0], list_data_var))

                        check_mid = [x for x in list_data_var if (x != data_test_design['Enum'][value_find_type_enum][-1] and x != data_test_design['Enum'][value_find_type_enum][0])]

                        tolerance = obj_level_2['Tolerance']
                        if __header__ == "IMPORTED PARAMETERS":
                            tolerance = 0
                        
                        if obj_level_2['Max'] > obj_level_2['Min']:
                            if ((obj_level_2['Max'] - tolerance) != obj_level_2['Min']):
                                if not (len(check_mid) != 0 and utils.checkMidValue(lst=check_mid, data=data_test_design['Enum'][value_find_type_enum], isEnum=True, min=obj_level_2['Min'], max=obj_level_2['Max'])):
                                    flag = False
                                    logger.error(">> Failed: Header {}, Variable {}, Type {}: Missing mid {}".format(__header__, key, value_find_type_enum, check_mid))
                                else:
                                    logger.info(">> Passed: Header {}, Variable {}, Type {}: Missing mid {}".format(__header__, key, value_find_type_enum, check_mid))
                            else:
                                if not (len(check_mid) == 0):
                                    flag = False
                                    logger.error(">> Failed: Header {}, Variable {}, Type {}: None Mid because MAX = MIN".format(__header__, key, value_find_type_enum))
                                else:
                                    logger.info(">> Passed: Header {}, Variable {}, Type {}: Missing mid {}".format(__header__, key, value_find_type_enum, check_mid))
                        elif obj_level_2['Max'] == obj_level_2['Min']:
                            if not (len(check_mid) == 0):
                                flag = False
                                logger.error(">> Failed: Header {}, Variable {}, Type {}: May be constant {}".format(__header__, key, value_find_type_enum, obj_level_2['Max']))
                            else:
                                logger.info(">> Passed: Header {}, Variable {}, Type {}: May be constant {}".format(__header__, key, value_find_type_enum, obj_level_2['Max']))
                        else:
                            flag = False
                            logger.error(">> Failed: Header {}, Variable {}, Type {}: BUG MAX NEVER LESS THAN MIN".format(__header__, key, obj_level_2['Type']))
                    else:
                        flag = False
                        logger.error(">> Failed: Header {}, Variable {}, Type {}: BUG".format(__header__, key, obj_level_2['Type']))
    except Exception as e:
        flag = False
        logger.exception(e)
    finally:
        logger.debug("Done")
        return flag

# Check information between summary xlsx, and test_summay_html is same or not
def check_information(file_test_summary_html, data, function_with_prj_name="", file_test_report_xml="", file_tpa="", file_CoverageReasonXLS="", opt=[], mode=""):
    def isTimeFormat(date):
        try:
            datetime.datetime.strptime(date, "%Y-%m-%dT%H:%M:%SZ")
            return True
        except ValueError:
            return False

    flag = False

    logger.debug("Check information {}", data.get("ItemName").replace(".c", ""))

    try:
        data_test_summary = FileTestSummaryHTML(file_test_summary_html).get_data()
        
        count = 0
        flag = True

        sub_new_func = ""
        if (utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM"):
            sub_new_func = function_with_prj_name
        else:
            sub_new_func = data.get("ItemName").replace(".c", "")

        if not (data_test_summary.get("Project") == sub_new_func):
            flag = False
            logger.error(">> Failed: Item FileTestSummaryHTML {} has different name: {} - {}".format(sub_new_func, data_test_summary.get("Project"), sub_new_func))
        else:
            logger.info(">> Passed: Item FileTestSummaryHTML {} has different name: {} - {}".format(sub_new_func, data_test_summary.get("Project"), sub_new_func))

        if not (data_test_summary.get("Verdict") == "Pass"):
            flag = False
            logger.error(">> Failed: Item FileTestSummaryHTML {} got different Verdict: {} - {}".format(data_test_summary.get("Project"), data_test_summary.get("Verdict"), data.get("Status Result")))
        else:
            logger.info(">> Passed: Item FileTestSummaryHTML {} got different Verdict: {} - {}".format(data_test_summary.get("Project"), data_test_summary.get("Verdict"), data.get("Status Result")))

        score_c0 = score_c1 = score_mcdc = ""

        if (data.get("C0") is None or data.get("C1") is None or data.get("MCDC") is None):
            flag = False
            logger.error("Item FileSummaryXLSX {} has none C0/C1/MCDC: {} - {} - {}".format(data.get("ItemName").replace(".c", ""), data.get("C0"), data.get("C1"), data.get("MCDC")))
        else:
            score_c0 = convert_score_percentage(data.get("C0"))
            score_c1 = convert_score_percentage(data.get("C1"))
            score_mcdc = convert_score_percentage(data.get("MCDC"))

            if not (check_score(score_test_summary=data_test_summary.get("C0"), score_exel=data.get("C0")) \
                    and check_score(score_test_summary=data_test_summary.get("C1"), score_exel=data.get("C1")) \
                    and check_score(score_test_summary=data_test_summary.get("MCDCU"), score_exel=data.get("MCDC"))):
                flag = False
                logger.error(">> Failed: Item FileTestSummaryHTML {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data_test_summary.get("Project"), data_test_summary.get("C0"), score_c0,
                                                                            data_test_summary.get("C1"), score_c1,
                                                                            data_test_summary.get("MCDCU"), score_mcdc)
                            )
            else:
                logger.info(">> Passed: Item FileTestSummaryHTML {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data_test_summary.get("Project"), data_test_summary.get("C0"), score_c0,
                                                                            data_test_summary.get("C1"), score_c1,
                                                                            data_test_summary.get("MCDCU"), score_mcdc)
                            )

        # Check information between FileTPA and Summary
        if(file_tpa != ""):
            try:
                data_tpa = FileTPA(file_tpa).get_data()
            except Exception as e:
                data_tpa = None
                logger.error(">> Failed: Item {} can not open file TPA. Please check at link: {}".format(data.get("ItemName").replace(".c", "") + ".c", file_tpa))

            if data_tpa is not None:
                data_tpa = data_tpa[data.get("ItemName").replace(".c", "") + ".c"]
                if (utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM"):
                    if not (check_score(score_test_summary=data_tpa.get("C0"), score_exel=data.get("C0")) \
                            and check_score(score_test_summary=data_tpa.get("C1"), score_exel=data.get("C1")) \
                            and check_score(score_test_summary=data_tpa.get("MC/DC"), score_exel=data.get("MCDC"))):
                        flag = False
                        logger.error(">> Failed: Item FileTPA {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data_tpa.get("FileName").replace(".c", ""), data_test_summary.get("C0"), score_c0,
                                                                                    data_tpa.get("C1"), score_c1,
                                                                                    data_tpa.get("MC/DC"), score_mcdc)
                                    )
                    else:
                        logger.info(">> Passed: Item FileTPA {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data_tpa.get("FileName").replace(".c", ""), data_test_summary.get("C0"), score_c0,
                                                                                    data_tpa.get("C1"), score_c1,
                                                                                    data_tpa.get("MC/DC"), score_mcdc)
                                    )
                else:
                    if not (check_score(score_test_summary=data_tpa.get("C0"), score_exel=data.get("C0")) \
                            and check_score(score_test_summary=data_tpa.get("C1"), score_exel=data.get("C1"))):
                        flag = False
                        logger.error(">> Failed: Item FileTPA {} has different C0: {}/{}; C1: {}/{}".format(data_tpa.get("FileName").replace(".c", ""), data_test_summary.get("C0"), score_c0,
                                                                                    data_tpa.get("C1"), score_c1)
                                    )
                    else:
                        logger.info(">> Passed: Item FileTPA {} has different C0: {}/{}; C1: {}/{}".format(data_tpa.get("FileName").replace(".c", ""), data_test_summary.get("C0"), score_c0,
                                                                                    data_tpa.get("C1"), score_c1)
                                    )
                
                if not (isTimeFormat(data_tpa.get("ExecutionDate"))):
                    flag = False
                    logger.error(">> Failed: Item FileTPA {} has different time format %Y-%m-%dT%H:%M:%SZ: {}".format(data_tpa.get("FileName").replace(".c", ""), data_tpa.get("ExecutionDate")))
                else:
                    logger.info(">> Passed: Item FileTPA {} has different time format %Y-%m-%dT%H:%M:%SZ: {}".format(data_tpa.get("FileName").replace(".c", ""), data_tpa.get("ExecutionDate")))
            
        if (utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM"):
            # Check information between FileTestReportXML and Summary
            data_test_report_xml = FileTestReportXML(file_test_report_xml).get_data()

            if not (check_score(score_test_summary=data_test_report_xml.get("C0"), score_exel=data.get("C0")) \
                    and check_score(score_test_summary=data_test_report_xml.get("C1"), score_exel=data.get("C1")) \
                    and check_score(score_test_summary=data_test_report_xml.get("MCDCU"), score_exel=data.get("MCDC"))):
                flag = False
                logger.error(">> Failed: Item FileTestReportXML {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data_test_report_xml.get("testScriptName"), data_test_report_xml.get("C0"), score_c0,
                                                                            data_test_report_xml.get("C1"), score_c1,
                                                                            data_test_report_xml.get("MCDCU"), score_mcdc)
                            )
            else:
                logger.info(">> Passed: Item FileTestReportXML {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data_test_report_xml.get("testScriptName"), data_test_report_xml.get("C0"), score_c0,
                                                                            data_test_report_xml.get("C1"), score_c1,
                                                                            data_test_report_xml.get("MCDCU"), score_mcdc)
                            )

            if ("check_FileCoverageReasonXLS" in opt and utils.load(CONST.SETTING, "mode_check_by_user").get("check_FileCoverageReasonXLS") == True):
                # Check information between FileCoverageReasonXLS and Summary
                data_CoverageReasonXLS = FileCoverageReasonXLS(file_CoverageReasonXLS).get_data()
                if not (check_score(score_test_summary=data_CoverageReasonXLS.get("C0"), score_exel=data.get("C0")) \
                        and check_score(score_test_summary=data_CoverageReasonXLS.get("C1"), score_exel=data.get("C1")) \
                        and check_score(score_test_summary=data_CoverageReasonXLS.get("MCDC"), score_exel=data.get("MCDC"))):
                    flag = False
                    logger.error(">> Failed: Item FileCoverageReasonXLS {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data_CoverageReasonXLS.get("Item_Name").replace(".c", ""), data_CoverageReasonXLS.get("C0"), score_c0,
                                                                                data_CoverageReasonXLS.get("C1"), score_c1,
                                                                                data_CoverageReasonXLS.get("MCDC"), score_mcdc)
                                )
                else:
                    logger.info(">> Passed: Item FileCoverageReasonXLS {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data_CoverageReasonXLS.get("Item_Name").replace(".c", ""), data_CoverageReasonXLS.get("C0"), score_c0,
                                                                                data_CoverageReasonXLS.get("C1"), score_c1,
                                                                                data_CoverageReasonXLS.get("MCDC"), score_mcdc)
                                )

        elif (utils.load(CONST.SETTING).get("sheetname") == "Merged_COEM"):
            if ("check_FileWalkThroughDoc" in opt):
                file_Walkthrough = utils.scan_files(Path(file_test_summary_html).parent.as_posix(), ext='.docx')[0][0]
                data_Walkthrough = FileWalkThroughDoc(file_Walkthrough).get_data()

                if not (data_Walkthrough.get("project").replace(".c", "") == sub_new_func and os.path.basename(file_Walkthrough) == "Walkthrough_Protocol_{}.docx".format(data.get("ItemName").replace(".c", ""))):
                    flag = False
                    logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong name - {}".format(data_Walkthrough.get("project").replace(".c", ""), sub_new_func))
                else:
                    logger.info(">> Passed: Item FileWalkThroughDoc {} has wrong name - {}".format(data_Walkthrough.get("project").replace(".c", ""), sub_new_func))

                if not (check_score(score_test_summary=data_Walkthrough.get("C0"), score_exel=data.get("C0")) \
                        and check_score(score_test_summary=data_Walkthrough.get("C1"), score_exel=data.get("C1"))):
                    flag = False
                    logger.error(">> Failed: Item FileWalkThroughDoc {} has different C0: {}/{}; C1: {}/{}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("C0"), score_c0,
                                                                                data_Walkthrough.get("C1"), score_c1))
                else:
                    logger.info(">> Passed: Item FileWalkThroughDoc {} has different C0: {}/{}; C1: {}/{}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("C0"), score_c0,
                                                                                data_Walkthrough.get("C1"), score_c1))

                if data.get("Baseline") == "" or data.get("Baseline") == "None" or data.get("Baseline") == None:
                    temp_baseline = ""
                else:
                    temp_baseline = data.get("Baseline")

                if not ((data_Walkthrough.get("baseline") == temp_baseline) or (temp_baseline == "" and data_Walkthrough.get("baseline") == "None")):
                    flag = False
                    logger.error(">> Failed: Item FileWalkThroughDoc {} has different Baseline: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("baseline"), temp_baseline))
                else:
                    logger.info(">> Passed: Item FileWalkThroughDoc {} has different Baseline: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("baseline"), temp_baseline))

                if not (data_Walkthrough.get("review partner") == convert_name(key=utils.load(CONST.SETTING, "users").get(data.get("Tester")).get("reviewer"), opt="name") \
                    and data_Walkthrough.get("review initiator") == convert_name(key=data.get("Tester"), opt="name")):
                    flag = False

                    logger.error(">> Failed: Item FileWalkThroughDoc has different reviewer/tester {}/{} - {}/{}".format(data_Walkthrough.get("review partner"), convert_name(key=utils.load(CONST.SETTING, "users").get(data.get("Tester")).get("reviewer"), opt="name"), \
                                                                                  data_Walkthrough.get("review initiator"), convert_name(key=data.get("Tester"), opt="name"))
                                )
                else:
                    logger.info(">> Passed: Item FileWalkThroughDoc has different reviewer/tester {}/{} - {}/{}".format(data_Walkthrough.get("review partner"), convert_name(key=utils.load(CONST.SETTING, "users").get(data.get("Tester")).get("reviewer"), opt="name"), \
                                                                                  data_Walkthrough.get("review initiator"), convert_name(key=data.get("Tester"), opt="name"))
                                )

                temp_path_test_WT = str(trim_src(data.get("ComponentName"))) + "\\Unit_tst\\" + str(data.get("TaskID")) + "\\" + data.get("ItemName").replace(".c", "")
                path_testscript = temp_path_test_WT + "\\Test_Spec"
                path_test_summary = temp_path_test_WT + "\\Test_Result"

                if str(data.get("TaskID")) == "9.1":
                    logger.warning("Item FileWalkThroughDoc {} has path: {} - {}. Please make sure it is correct".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get('path_testscript'),
                                                                                data_Walkthrough.get('path_test_summary'))
                                )
                else:
                    if not(data_Walkthrough.get('path_testscript') == path_testscript and data_Walkthrough.get('path_test_summary') == path_test_summary):
                        flag = False
                        logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong path: {}/{} - {}/{}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get('path_testscript'), path_testscript,
                                                                                    data_Walkthrough.get('path_test_summary'), path_test_summary)
                                    )
                    else:
                        logger.info(">> Passed: Item FileWalkThroughDoc {} has wrong path: {}/{} - {}/{}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get('path_testscript'), path_testscript,
                                                                                    data_Walkthrough.get('path_test_summary'), path_test_summary)
                                    )

                if (data.get("OPL/Defect") == "OPL" or data.get("OPL/Defect") == "Defect"):
                    if utils.load(CONST.SETTING, "mode_check_by_user").get("check_OPL_Walkthrough") == True:
                        num_OPL = check_OPL_Walkthrough(file_Walkthrough)

                        if not (num_OPL > 0):
                            flag = False
                            logger.error(">> Failed: Item FileWalkThroughDoc {} has none OPL: {}".format(data_Walkthrough.get("project").replace(".c", ""), str(num_OPL)))
                        else:
                            logger.info(">> Passed: Item FileWalkThroughDoc {} has none OPL: {}".format(data_Walkthrough.get("project").replace(".c", ""), str(num_OPL)))

                    if not (data_Walkthrough.get("tbl_finding").get('finding') != "/" \
                        and data_Walkthrough.get("tbl_finding").get('impact') != "/"):
                        flag = False
                        logger.error(">> Failed: Item FileWalkThroughDoc {} has none comment finding/impact: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('finding'),
                                                                                    data_Walkthrough.get("tbl_finding").get('impact'))
                                    )
                    else:
                        logger.info(">> Passed: Item FileWalkThroughDoc {} has none comment finding/impact: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('finding'),
                                                                                    data_Walkthrough.get("tbl_finding").get('impact'))
                                    )

                    if (data_Walkthrough.get('C0') == "100" and data_Walkthrough.get('C1') == "100"):
                        if not (data_Walkthrough.get("tbl_finding").get('confirm_UT9').strip() == "Yes, Documented"):
                            flag = False
                            logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong comment confirm UT9: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT9'),
                                                                                        '"Yes, Documented"')
                                        )
                        else:
                            logger.info(">> Passed: Item FileWalkThroughDoc {} has wrong comment confirm UT9: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT9'),
                                                                                        '"Yes, Documented"')
                                        )
                    else:
                        if not (data_Walkthrough.get("tbl_finding").get('confirm_UT9').strip() == "No, Documented"):
                            flag = False
                            logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong comment confirm UT9: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT9'),
                                                                                        '"No, Documented"')
                                    )
                        else:
                            logger.info(">> Passed: Item FileWalkThroughDoc {} has wrong comment confirm UT9: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT9'),
                                                                                        '"No, Documented"')
                                    ) 
                else:
                    if utils.load(CONST.SETTING, "mode_check_by_user").get("check_OPL_Walkthrough") == True:
                        num_OPL = check_OPL_Walkthrough(file_Walkthrough)

                        if not (num_OPL <= 0):
                            flag = False
                            logger.error(">> Failed: Item FileWalkThroughDoc {} has OPL: {}".format(data_Walkthrough.get("project").replace(".c", ""), str(num_OPL)))
                        else:
                            logger.info(">> Passed: Item FileWalkThroughDoc {} has OPL: {}".format(data_Walkthrough.get("project").replace(".c", ""), str(num_OPL)))

                    if not (data_Walkthrough.get("tbl_finding").get('finding') == "/" \
                        and data_Walkthrough.get("tbl_finding").get('impact') == "/"):
                        flag = False
                        logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong comment finding/impact: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('finding'),
                                                                                    data_Walkthrough.get("tbl_finding").get('impact'))
                                    )
                    else:
                        logger.info(">> Passed: Item FileWalkThroughDoc {} has wrong comment finding/impact: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('finding'),
                                                                                    data_Walkthrough.get("tbl_finding").get('impact'))
                                    )  

                    if not (data_Walkthrough.get("tbl_finding").get('confirm_UT9').strip() == "Yes"):
                        flag = False
                        logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong comment confirm UT9: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT9'),
                                                                                    '"Yes"')
                                    )
                    else:
                        logger.info(">> Passed: Item FileWalkThroughDoc {} has wrong comment confirm UT9: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT9'),
                                                                                    '"Yes"')
                                    )

        else:
            logger.error("No sheet name")

        # if (datetime.datetime.strptime(data_test_summary.get("date"), "%b %d, %Y, %H:%M %p").strftime("%d-%b-%Y") == datetime.datetime.strptime(data.get("End"), "%Y-%m-%d %H:%M:%S").strftime("%d-%b-%Y")):
        #     flag = True
        # else:
        #     flag = False
        #     logger.warning("ItemName {} got wrong date end: {} - {}".format(data_test_summary.get("Project"), datetime.datetime.strptime(data_test_summary.get("date"), "%b %d, %Y, %H:%M %p").strftime("%d-%b-%Y"), datetime.datetime.strptime(data.get("End"), "%Y-%m-%d %H:%M:%S").strftime("%d-%b-%Y")))
        #     return flag
    except Exception as e:
        flag = False
        logger.exception(e)
    finally:
        logger.debug("Done")
        return flag

def check_information_ASW(path, data, opt=""):
    def check_exist_plt(list_module, list_plt):
        flag = True
        for f in list_module:
            if not (f in list_plt):
                flag = False

        return flag
    
    flag = False
    try:
        logger.debug("Check information {}", data.get("ItemName").replace(".c", ""))

        flag = True

        lst_file_test_design = utils.scan_files(Path(path).as_posix(), ext='.xlsm')[0]
        lst_file_RTRT = utils.scan_files(Path(path).as_posix(), ext='ReportRTRT.txt')[0]
        lst_file_ATT = utils.scan_files(Path(path).as_posix(), ext='ATT_Report.xml')[0]
        lst_file_PLT = [os.path.basename(f).replace('.plt', '') for f in utils.scan_files(Path(path).as_posix(), ext='.plt')[0] if re.search('mdfFiles', str(Path(f).as_posix)) == None]
        lst_file_ATT_Prj = [f for f in utils.scan_files(Path(path).as_posix(), ext='.xml')[0] if re.search('{}/Project_.*.xml'.format(data.get("ItemName").replace(".c", "")), str(Path(f).as_posix))]

        lst_file_gif = []
        lst_file_htm = []

        if (utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM"):
            lst_file_test_design = [f for f in utils.scan_files(Path(path).as_posix(), ext='.xlsm')[0] if re.search('Documents/TD_.*.xlsm', str(Path(f).as_posix))]
            lst_file_RTRT = [f for f in utils.scan_files(Path(path).as_posix(), ext='ReportRTRT.txt')[0] if re.search('Delivery/TestResult', str(Path(f).as_posix))]
            lst_file_ATT = [f for f in utils.scan_files(Path(path).as_posix(), ext='ATT_Report.xml')[0] if re.search('Delivery/TestResult/.*/Rubasa/ATT_Report.xml', str(Path(f).as_posix))]
            lst_file_PLT = [os.path.basename(f).replace('.plt', '') for f in utils.scan_files(Path(path).as_posix(), ext='.plt')[0] if re.search('Documents/.*plt', str(Path(f).as_posix))]
            lst_file_ATT_Prj = [f for f in utils.scan_files(Path(path).as_posix(), ext='.xml')[0] if re.search('Delivery/TestEnvironment/.*/Project_.*.xml'.format(data.get("ItemName").replace(".c", "")), str(Path(f).as_posix))]
            lst_file_gif = [f for f in utils.scan_files(Path(path).as_posix(), ext='.xlsm')[0] if re.search('Documents/.*.gif', str(Path(f).as_posix))]
            lst_file_htm = [f for f in utils.scan_files(Path(path).as_posix(), ext='.xlsm')[0] if re.search('Documents/.*.htm', str(Path(f).as_posix))]

        file_test_design = file_RTRT = file_ATT = ""
        if not (len(lst_file_PLT)):
            flag = False
            logger.error(">> Failed: Item FilePLT {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_PLT))
        else:
            lst_file_PLT = lst_file_PLT
            logger.info(">> Passed: Item FilePLT {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_PLT))

        if not (len(lst_file_ATT_Prj) == 1):
            flag = False
            logger.error(">> Failed: Item file_ATT_Prj {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_ATT_Prj))
        else:
            logger.info(">> Passed: Item file_ATT_Prj {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_ATT_Prj))

        if (utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM"):
            if not (len(lst_file_gif) == 1):
                flag = False
                logger.error(">> Failed: Item File Gif {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_gif))
            else:
                logger.info(">> Passed: Item File Gif {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_gif))

            if not (len(lst_file_htm) == 1):
                flag = False
                logger.error(">> Failed: Item File HTM {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_htm))
            else:
                logger.info(">> Passed: Item File HTM {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_htm))

        data_test_design = data_RTRT = data_ATT = ""

        temp = ""
        if (utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM"):
            temp = data.get("ItemName").replace(".c", "")
        else:
            temp = data.get("ItemName").replace(".c", "")

        if not (len(lst_file_test_design) == 1):
            if len(utils.scan_files(Path(path).as_posix(), ext='.xls')[0]) or len(utils.scan_files(Path(path).as_posix(), ext='.xlsx')[0]) or len(utils.scan_files(Path(path).as_posix(), ext='.xlsb')[0]):
                flag = False
                logger.error(">> Failed: Item FileTestDesignXLSX {} has wrong extension. Please convert xls* to xlsm: {}".format(data.get("ItemName").replace(".c", ""), lst_file_test_design))
            else:
                flag = False
                logger.error(">> Failed: Item FileTestDesignXLSX {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_test_design))
        else:
            file_test_design = lst_file_test_design[0]
            data_test_design = FileTestDesignXLSX(file_test_design).get_data()

            logger.info(">> Passed: Item FileTestDesignXLSX {} has none file".format(data.get("ItemName").replace(".c", "")))
            if not (len(lst_file_ATT) == 1):
                flag = False
                logger.error(">> Failed: Item FileATTReportXML {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_ATT))
            else:
                logger.info(">> Passed: Item FileATTReportXML {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_ATT))

                file_ATT = lst_file_ATT[0]
                data_ATT = FileATTReportXML(file_ATT).get_data()

                if not (data_test_design.get("ItemName") == temp \
                    and data_ATT.get("ClassName") == temp):
                    flag = False
                    logger.error(">> Failed: Item FileTestDesignXLSX/FileATTReportXML has different name {} - {}".format(data_test_design.get("ItemName"), temp))
                else:
                    logger.info(">> Passed: Item FileTestDesignXLSX/FileATTReportXML has different name {} - {}".format(data_test_design.get("ItemName"), temp))

                for __each_data__ in data_test_design['DataTestDesign']:
                    if not (check_DataTestDesignXLSX(__each_data__)):
                        flag = False
                        logger.error(">> Failed: Item FileTestDesignXLSX {} with sheet {} has wrong information".format(data_test_design.get("ItemName"), __each_data__['sheetname']))
                    else:
                        logger.info(">> Passed: Item FileTestDesignXLSX {} with sheet {} has wrong information".format(data_test_design.get("ItemName"), __each_data__['sheetname']))

                if utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM":
                    if not(data_ATT.get("ClassVersion") == data.get("ItemRevision")):
                        flag = False
                        logger.error(">> Failed: Item FileTestDesignXLSX/FileATTReportXML {} got ItemRevision: {} - {}".format(data.get("ItemName"), data_ATT.get("ClassVersion"), data.get("ItemRevision")))
                    else:
                        logger.info(">> Passed: Item FileTestDesignXLSX/FileATTReportXML {} got ItemRevision: {} - {}".format(data.get("ItemName"), data_ATT.get("ClassVersion"), data.get("ItemRevision")))
                else:
                    if not(data_ATT.get("ClassVersion") == data.get("ItemRevision") and data_test_design.get("ItemRevision") == data.get("ItemRevision")):
                        flag = False
                        logger.error(">> Failed: Item FileTestDesignXLSX/FileATTReportXML {} got ItemRevision: {}/{} - {}".format(data.get("ItemName"), data_ATT.get("ClassVersion"), data_test_design.get("ItemRevision"), data.get("ItemRevision")))
                    else:
                        logger.info(">> Passed: Item FileTestDesignXLSX/FileATTReportXML {} got ItemRevision: {}/{} - {}".format(data.get("ItemName"), data_ATT.get("ClassVersion"), data_test_design.get("ItemRevision"), data.get("ItemRevision")))

                if not check_exist_plt(data_test_design.get("TM"), lst_file_PLT):
                    flag = False
                    logger.error(">> Failed: Item FileTestDesignXLSX/FilePLT {} has missing TestModuleName: {} - {}".format(data.get("ItemName"), data_test_design.get("TM"), lst_file_PLT))
                else:
                    logger.info(">> Passed: Item FileTestDesignXLSX/FilePLT {} has missing TestModuleName: {} - {}".format(data.get("ItemName"), data_ATT.get("TM"), lst_file_PLT))

                if not check_exist_plt(data_ATT.get("TestModuleName"), lst_file_PLT):
                    flag = False
                    logger.error(">> Failed: Item FileATTReportXML/FilePLT {} has missing TestModuleName: {} - {}".format(data.get("ItemName"), data_ATT.get("TestModuleName"), lst_file_PLT))
                else:
                    logger.info(">> Passed: Item FileATTReportXML/FilePLT {} has missing TestModuleName: {} - {}".format(data.get("ItemName"), data_ATT.get("TestModuleName"), lst_file_PLT))

                if not (data_ATT.get("CompleteVerdict") == "Passed"):
                    flag = False
                    logger.error(">> Failed: Item FileATTReportXML {} got Verdict: {} - {}".format(data.get("ItemName"), data_ATT.get("CompleteVerdict"), data.get("Status Result")))
                else:
                    logger.info(">> Passed: Item FileATTReportXML {} got Verdict: {} - {}".format(data.get("ItemName"), data_ATT.get("CompleteVerdict"), data.get("Status Result")))

        score_c0 = score_c1 = score_mcdc = ""

        if(data.get("C0") is None or data.get("C1") is None or data.get("MCDC") is None):
            flag = False
            logger.error("Item FileSummaryXLSX {} has none C0/C1/MCDC: {} - {} - {}".format(data.get("ItemName").replace(".c", ""), data.get("C0"), data.get("C1"), data.get("MCDC")))
        else:
            score_c0 = convert_score_percentage(data.get("C0"))
            score_c1 = convert_score_percentage(data.get("C1"))
            score_mcdc = convert_score_percentage(data.get("MCDC"))

        if not (len(lst_file_RTRT) == 1):
            if utils.load(CONST.SETTING).get("sheetname") == "Merged_JOEM":
                temp_path = Path(path).joinpath("Delivery", "TestResult", "yymmdd_hhmm_Sim_" + data.get("ItemName").replace(".c", "") + data.get("ItemRevision").replace(".", "_"), "CovRep_MCDC_" +  data.get("ItemName").replace(".c", "") + data.get("ItemRevision").replace(".", "_"), "ReportRTRT.txt")
                if(len(temp_path.as_posix()) > 256):
                    flag = False
                    logger.error(">> Failed: Item FileRTRTCov {} has long path".format(data.get("ItemName").replace(".c", "")))
                else:
                    flag = False
                    logger.error(">> Failed: Item FileRTRTCov {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_RTRT))
            else:
                flag = False
                logger.error(">> Failed: Item FileRTRTCov {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_RTRT))
        else:
            file_RTRT = lst_file_RTRT[0]
            data_RTRT = FileRTRTCov(file_RTRT).get_data()
            if not (check_score(score_test_summary=data_RTRT.get("C0"), score_exel=data.get("C0")) \
                    and check_score(score_test_summary=data_RTRT.get("C1"), score_exel=data.get("C1")) \
                    and check_score(score_test_summary=data_RTRT.get("MCDC"), score_exel=data.get("MCDC"))):
                flag = False
                logger.error(">> Failed: Item FileRTRTCov {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data.get("ItemName"), data_RTRT.get("C0"), score_c0,
                                                                            data_RTRT.get("C1"), score_c1,
                                                                            data_RTRT.get("MCDC"), score_mcdc)
                            )
            else:
                logger.info(">> Passed: Item FileRTRTCov {} has different C0: {}/{}; C1: {}/{}; MCDC: {}/{}".format(data.get("ItemName"), data_RTRT.get("C0"), score_c0,
                                                                            data_RTRT.get("C1"), score_c1,
                                                                            data_RTRT.get("MCDC"), score_mcdc)
                            )

        """ Check Walkthrough"""
        if (utils.load(CONST.SETTING).get("sheetname") == "Merged_COEM"):
            if ("check_FileWalkThroughDoc" in opt):
                lst_file_WT = [f for f in utils.scan_files(Path(path).as_posix(), ext='.docx')[0] if re.search('AR/{}/.*.docx'.format(data.get("ItemName").replace(".c", "")), str(Path(f).as_posix))]
                if not (len(lst_file_WT) == 1):
                    if len(utils.scan_files(Path(path).as_posix(), ext='.doc')[0]):
                        flag = False
                        logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong extension. Please convert doc* to docx: {}".format(data.get("ItemName").replace(".c", ""), lst_file_WT))
                    else:
                        flag = False
                        logger.error(">> Failed: Item FileWalkThroughDoc {} has none file: {}".format(data.get("ItemName").replace(".c", ""), lst_file_WT))
                else:
                    file_Walkthrough = lst_file_WT[0]
                    data_Walkthrough = FileWalkThroughDoc(file_Walkthrough).get_data(opt="ASW")

                    if not (data_Walkthrough.get("project").replace(".c", "") == data.get("ItemName") and os.path.basename(file_Walkthrough) == "Walkthrough_Protocol_{}.docx".format(data.get("ItemName").replace(".c", ""))):
                        flag = False
                        logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong name - {}".format(data_Walkthrough.get("project").replace(".c", ""), data.get("ItemName")))
                    else:
                        logger.info(">> Passed: Item FileWalkThroughDoc {} has wrong name - {}".format(data_Walkthrough.get("project").replace(".c", ""), data.get("ItemName")))

                    if not(data_Walkthrough.get("ItemRevision") == data.get("ItemRevision")):
                        flag = False
                        logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong ItemRevision: {} - {}".format(data.get("ItemName"), data_Walkthrough.get("ItemRevision"), data.get("ItemRevision")))
                    else:
                        logger.info(">> Passed: Item FileWalkThroughDoc {} has wrong ItemRevision: {} - {}".format(data.get("ItemName"), data_Walkthrough.get("ItemRevision"), data.get("ItemRevision")))

                    if not (check_score(score_test_summary=data_Walkthrough.get("C0"), score_exel=data.get("C0")) \
                            and check_score(score_test_summary=data_Walkthrough.get("C1"), score_exel=data.get("C1"))):
                        flag = False
                        logger.error(">> Failed: Item FileWalkThroughDoc {} has different C0: {}/{}; C1: {}/{}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("C0"), score_c0,
                                                                                    data_Walkthrough.get("C1"), score_c1))
                    else:
                        logger.info(">> Passed: Item FileWalkThroughDoc {} has different C0: {}/{}; C1: {}/{}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("C0"), score_c0,
                                                                                    data_Walkthrough.get("C1"), score_c1))

                    if data.get("Baseline") == "" or data.get("Baseline") == "None" or data.get("Baseline") == None:
                        temp_baseline = ""
                    else:
                        temp_baseline = data.get("Baseline")

                    if not ((data_Walkthrough.get("baseline") == temp_baseline) or (temp_baseline == "" and data_Walkthrough.get("baseline") == "None")):
                        flag = False
                        logger.error(">> Failed: Item FileWalkThroughDoc {} has different Baseline: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("baseline"), temp_baseline))
                    else:
                        logger.info(">> Passed: Item FileWalkThroughDoc {} has different Baseline: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("baseline"), temp_baseline))

                    if not (data_Walkthrough.get("review partner") == convert_name(key=utils.load(CONST.SETTING, "users").get(data.get("Tester")).get("reviewer"), opt="name") \
                        and data_Walkthrough.get("review initiator") == convert_name(key=data.get("Tester"), opt="name")):
                        flag = False
                        logger.error(">> Failed: Item FileWalkThroughDoc has different reviewer/tester {}/{} - {}/{}".format(data_Walkthrough.get("review partner"), convert_name(key=utils.load(CONST.SETTING, "users").get(data.get("Tester")).get("reviewer"), opt="name"), \
                                                                                    data_Walkthrough.get("review initiator"), convert_name(key=data.get("Tester"), opt="name"))
                                    )
                    else:
                        logger.info(">> Passed: Item FileWalkThroughDoc has different reviewer/tester {}/{} - {}/{}".format(data_Walkthrough.get("review partner"), convert_name(key=utils.load(CONST.SETTING, "users").get(data.get("Tester")).get("reviewer"), opt="name"), \
                                                                                    data_Walkthrough.get("review initiator"), convert_name(key=data.get("Tester"), opt="name"))
                                    )

                    if (data.get("OPL/Defect") == "OPL" or data.get("OPL/Defect") == "Defect"):
                        if utils.load(CONST.SETTING, "mode_check_by_user").get("check_OPL_Walkthrough") == True:
                            num_OPL = check_OPL_Walkthrough(file_Walkthrough)

                            if not (num_OPL > 0):
                                flag = False
                                logger.error(">> Failed: Item FileWalkThroughDoc {} has none OPL: {}".format(data_Walkthrough.get("project").replace(".c", ""), str(num_OPL)))
                            else:
                                logger.info(">> Passed: Item FileWalkThroughDoc {} has none OPL: {}".format(data_Walkthrough.get("project").replace(".c", ""), str(num_OPL)))

                        if not (data_Walkthrough.get("tbl_finding").get('finding') != "/" \
                            and data_Walkthrough.get("tbl_finding").get('impact') != "/"):
                            flag = False
                            logger.error(">> Failed: Item FileWalkThroughDoc {} has none comment finding/impact: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('finding'),
                                                                                        data_Walkthrough.get("tbl_finding").get('impact'))
                                        )
                        else:
                            logger.info(">> Passed: Item FileWalkThroughDoc {} has none comment finding/impact: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('finding'),
                                                                                        data_Walkthrough.get("tbl_finding").get('impact'))
                                        )

                        if (data_Walkthrough.get('C0') == "100" and data_Walkthrough.get('C1') == "100"):
                            if not (data_Walkthrough.get("tbl_finding").get('confirm_UT26').strip() == "Yes, Documented"):
                                flag = False
                                logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong comment confirm UT26: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT26'), '"Yes, Documented"'))
                            else:
                                logger.info(">> Passed: Item FileWalkThroughDoc {} has wrong comment confirm UT26: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT26'), '"Yes, Documented"'))
                        else:
                            if not (data_Walkthrough.get("tbl_finding").get('confirm_UT26').strip() == "No, Documented"):
                                flag = False
                                logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong comment confirm UT26: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT26'), '"No, Documented"'))
                            else:
                                logger.info(">> Passed: Item FileWalkThroughDoc {} has wrong comment confirm UT26: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT26'), '"No, Documented"'))
                    else:
                        if utils.load(CONST.SETTING, "mode_check_by_user").get("check_OPL_Walkthrough") == True:
                            num_OPL = check_OPL_Walkthrough(file_Walkthrough)

                            if not (num_OPL <= 0):
                                flag = False
                                logger.error(">> Failed: Item FileWalkThroughDoc {} has OPL: {}".format(data_Walkthrough.get("project").replace(".c", ""), str(num_OPL)))
                            else:
                                logger.info(">> Passed: Item FileWalkThroughDoc {} has OPL: {}".format(data_Walkthrough.get("project").replace(".c", ""), str(num_OPL)))

                        if not (data_Walkthrough.get("tbl_finding").get('finding') == "/" \
                            and data_Walkthrough.get("tbl_finding").get('impact') == "/"):
                            flag = False
                            logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong comment finding/impact: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('finding'),
                                                                                        data_Walkthrough.get("tbl_finding").get('impact'))
                                        )
                        else:
                            logger.info(">> Passed: Item FileWalkThroughDoc {} has wrong comment finding/impact: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('finding'),
                                                                                        data_Walkthrough.get("tbl_finding").get('impact'))
                                        )

                        if not (data_Walkthrough.get("tbl_finding").get('confirm_UT26').strip() == ""):
                            flag = False
                            logger.error(">> Failed: Item FileWalkThroughDoc {} has wrong comment confirm UT26: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT26'), '""'))
                        else:
                            logger.info(">> Passed: Item FileWalkThroughDoc {} has wrong comment confirm UT26: {} - {}".format(data_Walkthrough.get("project").replace(".c", ""), data_Walkthrough.get("tbl_finding").get('confirm_UT26'), '""'))

    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")
        return flag

# Check Release for JOEM is correct ot not
def check_archives_joem(path_summary, dir_input, taskids, begin=59, end=59):
    logger.debug("Start checker: Archives")
    flag = False
    try:
        doc = FileSummaryXLSX(path_summary)
        data = doc.parse2json(begin=begin, end=end)

        file_log = open("log_delivery.txt", "w")

        print("Start checker: Archives")
        print("*****************************************************************")
        file_log.write("Start checker: Archives\n")
        file_log.write("*****************************************************************\n")

        for taskid in taskids["TaskGroup"]:
            temp_data_prj = doc.get_data(data=data, key="Project", value=taskids["Project"])
            data_taskid = doc.get_data(data=temp_data_prj, key="TaskGroup", value=taskid)
            
            path_taskid = Path(dir_input).joinpath(str(taskid))
            if (path_taskid.exists()):
                count = 0
                for item in data_taskid.keys():

                    function = data_taskid[item].get("ItemName").replace(".c", "")
                    user_tester = data_taskid[item].get("Tester")
                    mt_number = data_taskid[item].get("MT_Number").replace("UT_", "").replace("MT_", "")
                    bb_number = taskids["BB"]

                    if "ASW" == data_taskid[item].get("Type"):
                        mt_number = "MT_" + mt_number
                        bb_number = ""
                    elif "PSW" == data_taskid[item].get("Type"):
                        mt_number = "UT_" + mt_number
                        bb_number = "_" + bb_number
                    else:
                        mt_number = "NONE"
                        logger.error("BUG mt_number")

                    folder_mt_function = "{}_{}{}".format(mt_number, function, bb_number)

                    b_check_exist = check_exist(dir_input=path_taskid, function=folder_mt_function)
                    if (b_check_exist):
                        count += 1

                        if "ASW" == data_taskid[item].get("Type"):
                            if check_information_ASW(path=Path.joinpath(path_taskid, folder_mt_function), data=data_taskid[item]):
                                print("{},{},{},{},{}".format(taskid, mt_number, function, user_tester, "OK"))
                                file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "OK"))
                            else:
                                logger.error("Different Information {},{},{},{},{}".format(taskid, mt_number, function, user_tester, "NG_DiffInfor"))
                                file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "NG_DiffInfor"))

                        elif "PSW" == data_taskid[item].get("Type"):
                            file_tpa = Path(path_taskid).joinpath(folder_mt_function, "Cantata", "results", "{}.tpa".format(function))
                            file_CoverageReasonXLS = Path(path_taskid).joinpath(folder_mt_function, "doc", "{}_{}".format(function, "CodeCoverage_or_Fail_Reason.xls"))
                            file_test_report_xml = Path(path_taskid).joinpath(folder_mt_function, "Cantata", "results", "test_report.xml")
                            file_test_summary = Path(path_taskid).joinpath(folder_mt_function, "Cantata", "results", "test_summary.html")

                            option_check = []
                            # if (len(str(function)) < 32):
                            #     option_check.append("check_FileCoverageReasonXLS")
                            # else:
                            #     option_check = []
                            option_check.append("check_FileCoverageReasonXLS")

                            if check_information(file_test_summary_html=file_test_summary, data=data_taskid[item], function_with_prj_name=folder_mt_function, file_test_report_xml=file_test_report_xml, file_tpa=file_tpa, file_CoverageReasonXLS=file_CoverageReasonXLS, opt=option_check):
                                print("{},{},{},{},{}".format(taskid, mt_number, function, user_tester, "OK"))
                                file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "OK"))
                            else:
                                logger.error("Different Information {},{},{},{},{}".format(taskid, mt_number, function, user_tester, "NG_DiffInfor"))
                                file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "NG_DiffInfor"))
                        else:
                            logger.error("Bug No Type")

                    else:
                        logging.warning("{},{},{},{},{}".format(taskid, mt_number, function, user_tester, "NG"))
                        file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "NG"))

                num_tpa = len(utils.scan_files(directory=path_taskid, ext=".tpa")[0]) + len(utils.scan_files(directory=path_taskid, ext=".xlsm")[0])
                status = ["GOOD" if (num_tpa == len(data_taskid)) and (count == len(data_taskid)) else "BAD"][0]
                print("## Total {}: status {}: Found/Count/Total - {}/{}/{}".format(str(taskid), status, count, num_tpa, len(data_taskid)))
                print("-----------------------------------------------------------------\n")

                file_log.write("## Total {}: status {}: Found/Count/Total - {}/{}/{}\n".format(str(taskid), status, count, num_tpa, len(data_taskid)))
                file_log.write("-----------------------------------------------------------------\n")

            else:
                logger.warning("TaskID {} is not existed".format(path_taskid))
                file_log.write("TaskID {} is not existed\n".format(path_taskid))
                next

        print("FINISH")
        file_log.write("FINISH\n")
        file_log.close()

    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Check Release for JOEM is correct ot not
def make_archives_joem(path_summary, dir_input, dir_output, taskids, begin=59, end=59):
    logger.debug("Start checker: Make Archives JOEM")
    try:
        doc = FileSummaryXLSX(path_summary)
        data = doc.parse2json(begin=begin, end=end)

        file_log = open("log_delivery.txt", "w")

        print("Start checker: Make Archives JOEM")
        print("*****************************************************************")
        file_log.write("Start checker: Make Archives JOEM\n")
        file_log.write("*****************************************************************\n")

        for taskid in taskids["TaskGroup"]:
            temp_data_prj = doc.get_data(data=data, key="Project", value=taskids["Project"])
            data_taskid = doc.get_data(data=temp_data_prj, key="TaskGroup", value=taskid)
            bb_number = taskids["BB"]
            path_taskid = Path(dir_input).joinpath(str(taskid))
            if (path_taskid.exists()):
                count = 0
                for item in data_taskid.keys():
                    function = data_taskid[item].get("ItemName").replace(".c", "")
                    user_tester = data_taskid[item].get("Tester")
                    mt_number = data_taskid[item].get("MT_Number").replace("UT_", "").replace("MT_", "")

                    if "ASW" == data_taskid[item].get("Type"):
                        mt_number = "MT_" + mt_number
                        logging.warning("{},{},{},{},{}".format(taskid, mt_number, function, user_tester, "NG_MT_Check_Later"))
                        file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "NG_MT_Check_Later"))
                        continue
                    elif "PSW" == data_taskid[item].get("Type"):
                        mt_number = "UT_" + mt_number
                    else:
                        mt_number = "NONE"
                        print("BUG mt_number")

                    folder_mt_function = "{}_{}_{}".format(mt_number, function, bb_number)

                    b_check_exist = check_exist(dir_input=path_taskid, function=folder_mt_function)
                    if (b_check_exist):
                        count += 1

                        file_tpa = Path(path_taskid).joinpath(folder_mt_function, "Cantata", "results", "{}.tpa".format(function))
                        file_CoverageReasonXLS = Path(path_taskid).joinpath(folder_mt_function, "doc", "{}_{}".format(function, "CodeCoverage_or_Fail_Reason.xls"))
                        file_test_report_xml = Path(path_taskid).joinpath(folder_mt_function, "Cantata", "results", "test_report.xml")
                        file_test_summary = Path(path_taskid).joinpath(folder_mt_function, "Cantata", "results", "test_summary.html")

                        if check_information(file_test_summary_html=file_test_summary, data=data_taskid[item], function_with_prj_name=folder_mt_function, file_test_report_xml=file_test_report_xml, opt=[]):
                            utils.copy(src=Path(CONST.TEMPLATE).joinpath("JOEM", "template.tpa"), dst=file_tpa)
                            update_tpa(file=file_tpa, data=data_taskid[item], file_test_summary_html=file_test_summary)
                            try:
                                FileCoverageReasonXLS(file_CoverageReasonXLS).update(data_taskid[item])
                                print("{},{},{},{}".format(taskid, function, user_tester, "OK"))
                                file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "OK"))
                            except Exception as e:
                                logger.error("{},{},{},{}".format(taskid, function, user_tester, "NG_FileCoverageReasonXLS_Long_Name"))
                                file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG_Long_Name"))
                        else:
                            logger.error("Different Information {},{},{},{},{}".format(taskid, mt_number, function, user_tester, "NG_DiffInfor"))
                            file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "NG_DiffInfor"))

                    else:
                        logging.warning("{},{},{},{},{}".format(taskid, mt_number, function, user_tester, "NG"))
                        file_log.write("{},{},{},{},{}\n".format(taskid, mt_number, function, user_tester, "NG"))

                num_tpa = len(utils.scan_files(directory=path_taskid, ext=".tpa")[0]) + len(utils.scan_files(directory=path_taskid, ext=".xlsm")[0])

                status = ["GOOD" if (num_tpa == len(data_taskid)) and (count == len(data_taskid)) else "BAD"][0]
                print("## Total {}: status {}: Found/Count/Total - {}/{}/{}".format(str(taskid), status, count, num_tpa, len(data_taskid)))
                print("-----------------------------------------------------------------\n")

                file_log.write("## Total {}: status {}: Found/Count/Total - {}/{}/{}\n".format(str(taskid), status, count, num_tpa, len(data_taskid)))
                file_log.write("-----------------------------------------------------------------\n")

            else:
                logger.warning("{} is not existed".format(path_taskid))
                file_log.write("{} is not existed\n".format(path_taskid))
                next

        print("FINISH")
        file_log.write("FINISH\n")
        file_log.close()

    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Check Release is correct or not
def check_releases(path_summary, dir_input, taskids, begin=59, end=59):
    logger.debug("Start checker: Release")

    try:
        doc = FileSummaryXLSX(path_summary)
        data = doc.parse2json(begin=begin, end=end)
        print("Start checker: Release")
        print("*****************************************************************")
        file_log = open("log_delivery.txt", "w")
        file_log.write("Start checker: Release\n")
        file_log.write("*****************************************************************\n")
        for taskid in taskids:
            data_taskid = doc.get_data(data=data, key="TaskID", value=taskid)
            path_taskid = Path(dir_input).joinpath(str(taskid), "RV")
            if (path_taskid.exists()):
                count = 0
                for item in data_taskid.keys():
                    function = data_taskid[item].get("ItemName").replace(".c", "")
                    user_tester = data_taskid[item].get("Tester")
                    b_check_exist = check_exist(dir_input=path_taskid, function=function)
                    if (b_check_exist):
                        count += 1
                        print("{},{},{},{}".format(taskid, function, user_tester, "OK"))
                        file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "OK"))

                    else:
                        logger.warning("{},{},{},{}".format(taskid, function, user_tester, "NG"))
                        file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG"))

                status = ["GOOD" if (len(os.listdir(path_taskid)) == len(data_taskid)) and (count == len(data_taskid)) else "BAD"][0]
                print("## Total {}: status {}: Found/Count/Total - {}/{}/{}".format(str(taskid), status, count, len(os.listdir(path_taskid)), len(data_taskid)))
                print("-----------------------------------------------------------------\n")

                file_log.write("## Total {}: status {}: Found/Count/Total - {}/{}/{}\n".format(str(taskid), status, count, len(os.listdir(path_taskid)), len(data_taskid)))
                file_log.write("-----------------------------------------------------------------\n")
            else:
                logger.error("{} is not existed".format(path_taskid))
                file_log.write("{} is not existed\n".format(path_taskid))
                next

        print("FINISH")
        file_log.write("FINISH\n")
        file_log.close()
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Check Archive is correct or not
def check_archives(path_summary, dir_input, taskids, begin=59, end=59):
    logger.debug("Start checker: Archives")
    try:
        doc = FileSummaryXLSX(path_summary)
        data = doc.parse2json(begin=begin, end=end)
        file_log = open("log_delivery.txt", "a")

        print("Start checker: Archives")
        print("*****************************************************************")
        file_log.write("Start checker: Archives\n")
        file_log.write("*****************************************************************\n")
        for taskid in taskids:
            data_taskid = doc.get_data(data=data, key="TaskID", value=taskid)
            path_taskid = Path(dir_input).joinpath(str(taskid), "AR")
            if (path_taskid.exists()):
                count = 0
                for item in data_taskid.keys():
                    function = data_taskid[item].get("ItemName").replace(".c", "")
                    user_tester = data_taskid[item].get("Tester")
                    path_with_component = ""
                    if "ASW" == data_taskid[item].get("Type"):
                        path_with_component = Path(path_taskid)
                    elif "PSW" == data_taskid[item].get("Type"):
                        path_with_component = Path(path_taskid).joinpath(str(trim_src(data_taskid[item].get("ComponentName"))), "Unit_tst", str(data_taskid[item].get("TaskID")))
                        if str(taskid) == "9.1":
                            path_with_component = Path(path_taskid)
                    else:
                        print("BUG No Type")

                    b_check_exist = check_exist(dir_input=path_with_component, function=function)
                    if (b_check_exist):
                        count += 1

                        if "ASW" == data_taskid[item].get("Type"):
                            if check_information_ASW(path=Path.joinpath(path_with_component, function), data=data_taskid[item], opt=["check_FileWalkThroughDoc"]):
                                print("{},{},{},{}".format(taskid, function, user_tester, "OK"))
                                file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "OK"))
                            else:
                                logger.error("Different Information {},{},{},{}".format(taskid, function, user_tester, "NG_DiffInfor"))
                                file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG_DiffInfor"))
                        elif "PSW" == data_taskid[item].get("Type"):
                            f_test_summary = Path(path_taskid).joinpath(path_with_component, function, "Test_Result", "test_summary.html")
                            f_tpa = Path(path_taskid).joinpath(path_with_component, function, "Test_Result","{}.tpa".format(function))

                            if str(taskid) == "9.1":
                                f_test_summary = Path(path_taskid).joinpath(path_with_component, function, "test_summary.html")
                                f_tpa = ""

                            if check_information(file_test_summary_html=f_test_summary, file_tpa=f_tpa, data=data_taskid[item], opt=["check_FileWalkThroughDoc"]):
                                print("{},{},{},{}".format(taskid, function, user_tester, "OK"))
                                file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "OK"))
                            else:
                                logger.error("Different Information {},{},{},{}".format(taskid, function, user_tester, "NG_DiffInfor"))
                                file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG_DiffInfor"))
                        else:
                            logger.error("Bug: No type")

                    else:
                        logging.warning("{},{},{},{}".format(taskid, function, user_tester, "NG"))
                        file_log.write("{},{},{},{}\n".format(taskid, function, user_tester, "NG"))

                num_tpa = len(utils.scan_files(directory=path_taskid, ext=".tpa")[0])
                status = ["GOOD" if (num_tpa == len(data_taskid)) and (count == len(data_taskid)) else "BAD"][0]
                print("## Total {}: status {}: Found/Count/Total - {}/{}/{}".format(str(taskid), status, count, num_tpa, len(data_taskid)))
                print("-----------------------------------------------------------------\n")

                file_log.write("## Total {}: status {}: Found/Count/Total - {}/{}/{}\n".format(str(taskid), status, count, num_tpa, len(data_taskid)))
                file_log.write("-----------------------------------------------------------------\n")

            else:
                logger.warning("{} is not existed".format(path_taskid))
                file_log.write("{} is not existed\n".format(path_taskid))
                next

        print("FINISH")
        file_log.write("FINISH\n")
        file_log.close()

    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Update WalkThrough
def update_walkthrough(file, data, file_test_summary_html):
    logger.debug("Update Walkthrough {}", file)
    try:
        data_test_summary = FileTestSummaryHTML(file_test_summary_html).get_data()
        temp = str(trim_src(data.get("ComponentName"))) + "\\Unit_tst\\" + str(data.get("TaskID")) + "\\" + data.get("ItemName").replace(".c", "")

        data_baseline = data.get("Baseline")
        if data_baseline == None or data_baseline == "" or data_baseline == "None":
            data_baseline = ""

        reviewer = convert_name(key=utils.load(CONST.SETTING, "reviewer"), opt="name")

        dict_walkthrough = {
            'date': datetime.datetime.now().strftime("%m/%d/%Y"),
            'project': data.get("ItemName"),
            'review initiator': convert_name(key=data.get("Tester"), opt="name"),
            'effort': str(0.5),
            'baseline': data_baseline,
            'review partner' : reviewer,
            'path_testscript': temp + "\\Test_Spec",
            'path_test_summary': temp + "\\Test_Result",
            'ScoreC0C1': " Test summary\n\tC0: " + data_test_summary.get("C0") + "%\tC1: " + data_test_summary.get("C1") + "%",
        }

        FileWalkThroughDoc(file).update(dict_walkthrough)

    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Update File TPA
def update_tpa(file, data, file_test_summary_html):
    logger.debug("Update tpa {}", file)

    try:
        data_test_summary = FileTestSummaryHTML(file_test_summary_html).get_data()
        data_tpa = {
            "UnitUnderTest": data.get("ItemName").replace(".c", "") + ".c",
            "NTUserID": str(convert_name(key=data.get("Tester"), opt="id")),
            "ExecutionDate" : datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"),
            "FileName": data.get("ItemName"),
            "Verdict": data_test_summary.get('Verdict').replace("Pass", "Passed").replace("Fail", "Failed"),
            "C0": data_test_summary.get('C0'),
            "C1": data_test_summary.get('C1'),
            "MCDCU": data_test_summary.get('MCDCU'),
        }

        FileTPA(file).update_tpa(data_tpa)
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Encrypt file with 7z
def sevenzip(filename, zipname):
    try:
        zip_exe_path = utils.load(CONST.SETTING, "Tool_7z")
        with open(os.devnull, 'w') as null:
            subprocess.call([zip_exe_path, 'a', '-tzip', zipname, filename], stdout=null, stderr=null)
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Create Archive Walkthrough
def make_archieves(path_summary, dir_input, dir_output, taskids, begin=59, end=59):
    logger.debug("Create Archive Walkthrough")
    try:

        doc = FileSummaryXLSX(path_summary)
        data = doc.parse2json(begin=begin, end=end)

        print("Start checker: Make Archives")
        print("*****************************************************************")

        for taskid in taskids:
            data_taskid = doc.get_data(data=data, key="TaskID", value=taskid)
            path_taskid = Path(dir_input).joinpath(str(taskid), "RV")
            if (path_taskid.exists()):
                count = 0
                for item in data_taskid.keys():
                    function = data_taskid[item].get("ItemName").replace(".c", "")
                    user_tester = data_taskid[item].get("Tester")

                    b_check_exist = check_exist(dir_input=path_taskid, function=function)
                    if (b_check_exist):
                        count += 1
                        if "ASW" == data_taskid[item].get("Type"):
                            logger.warning("{},{},{},{}".format(taskid, function, user_tester, "ASW_No_Need"))
                            continue
                        elif "PSW" == data_taskid[item].get("Type"):
                            temp = str(data_taskid[item].get("C0"))
                            if (temp != "-" or temp != ""):
                                temp_component = str(data_taskid[item].get("ComponentName"))
                                if is_have_src(temp_component):
                                    final_dst = Path(dir_output).joinpath(str(taskid), "AR", trim_src(temp_component), "Unit_tst", str(taskid), function)
                                    dir_Configuration = Path(final_dst).joinpath("Configuration")
                                    dir_Test_Spec = Path(final_dst).joinpath("Test_Spec")
                                    dir_Test_Result = Path(final_dst).joinpath("Test_Result")

                                    f_walkthrough = Path(dir_Test_Result).joinpath("Walkthrough_Protocol_" + function + ".docx")
                                    f_tpa = Path(dir_Test_Result).joinpath(function + ".tpa")
                                    f_test_summary = Path(path_taskid).joinpath(function, "Cantata", "results", "test_summary.html")

                                    if check_information(file_test_summary_html=f_test_summary, data=data_taskid[item]):
                                        Path(dir_Configuration).parent.mkdir(parents=True, exist_ok=True)
                                        Path(dir_Configuration).mkdir(exist_ok=True)
                                        Path(dir_Test_Spec).mkdir(exist_ok=True)
                                        Path(dir_Test_Result).mkdir(exist_ok=True)

                                        utils.copy(src=Path(CONST.TEMPLATE).joinpath("COEM", "WT_template_PSW.docx"), dst=f_walkthrough)
                                        utils.copy(src=Path(CONST.TEMPLATE).joinpath("COEM", "template.tpa"), dst=f_tpa)
                                        utils.copy(src=f_test_summary, dst=dir_Test_Result)

                                        update_walkthrough(file=f_walkthrough, data=data_taskid[item], file_test_summary_html=f_test_summary)
                                        update_tpa(file=f_tpa, data=data_taskid[item], file_test_summary_html=f_test_summary)

                                        sevenzip(filename=Path(path_taskid).joinpath(function).as_posix(), zipname=Path(dir_Configuration).joinpath(str(function) + ".zip").as_posix())

                                        for f in utils.scan_files(directory=Path(path_taskid).joinpath(function, "Cantata", "tests"), ext=".c")[0]:
                                            sevenzip(filename=f.as_posix(), zipname=Path(dir_Test_Spec).joinpath(os.path.basename(f).replace(".c", ".zip")).as_posix())

                                        print("{},{},{},{}".format(taskid, function, user_tester, "OK"))
                                    else:
                                        logger.error("Different information {},{},{},{}".format(taskid, function, user_tester, "NG"))
                                        next

                                else:
                                    logger.error("Miss src in componentname {},{},{},{}".format(taskid, function, user_tester, "NG"))
                        else:
                            logger.error("BUG No Type")
                    else:
                        logger.warning("{},{},{},{}".format(taskid, function, user_tester, "NG"))

                status = ["GOOD" if (len(os.listdir(path_taskid)) == len(data_taskid)) and (count == len(data_taskid)) else "BAD"][0]
                print("## Total {}: status {}: Found/Count/Total - {}/{}/{}".format(str(taskid), status, count, len(os.listdir(path_taskid)), len(data_taskid)))
                print("-----------------------------------------------------------------\n")
            else:
                next
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

def create_summary_json_file(file_summary, sheetname="", begin=59, end=59):
    # Generate json file
    if sheetname == "":
        sheetname = utils.load(CONST.SETTING).get("sheetname")

    file_log = Path(__file__).parent.joinpath("log_json", "{}_{}_{}.json".format("log_summary", sheetname, datetime.datetime.now().strftime("%Y_%m_%dT%H_%MZ"), ".json"))
    with open(file_log, errors='ignore', mode='w') as fp:
        json.dump(FileSummaryXLSX(file_summary).parse2json(sheetname=sheetname, begin=begin, end=end), fp, indent=4, sort_keys=True)

def collect_information_deliverables(file_summary, sheetname="", begin=59, end=59):
    if sheetname == "":
        sheetname = utils.load(CONST.SETTING).get("sheetname")

    print("Start checker: Make collect_information_deliverables for: {}".format(sheetname))
    print("*****************************************************************")
    data = FileSummaryXLSX(file_summary).parse2json(sheetname=sheetname, begin=begin, end=end)

    lst_header = ["ELOC Recheck With Tool", "LOC Complete", "Planned Start", "Planned End", "Release Date", "OPL/Defect"]
    result_header = ["Project", "Type", "Assigned task (ELOC)", "Assigned date", "Target date", "Delivered task (ELOC)", "Delivered date", "Remain (ELOC)", "% Completion", "Nums of Defect"]

    template_json_obj = {
        "Project": None,
        "Type": ["ASW", "PSW"],
        "data": {
            "Assigned task (ELOC)": {"ASW": 0, "PSW": 0},
            "Assigned date": {"ASW": list(), "PSW": list()},
            "Target date": {"ASW": list(), "PSW": list()},
            "Delivered task (ELOC)": {"ASW": 0, "PSW": 0},
            "Delivered date": {"ASW": list(), "PSW": list()},
            "Remain (ELOC)": {"ASW": -1, "PSW": -1},
            "% Completion": {"ASW": -1, "PSW": -1},
            "Nums of Defect": {"ASW": 0, "PSW": 0}
        }
    }

    item = -1
    l_prj = list()
    for item in data.keys():
        if data[item].get("Project") is not None:
            l_prj.append(data[item].get("Project"))

    l_prj = list(set(l_prj))
    l_prj.sort()

    item = -1

    result = dict()
    data_prj = dict()

    for index, prj in enumerate(l_prj):
        data_prj = copy.deepcopy(template_json_obj)
        for item in data.keys():
            if(data[item].get("Project") == prj):
                data_prj['Project'] = prj
                for __header__ in lst_header:
                    if data[item].get(__header__) is not None and (data[item].get(__header__) != "None"):
                        if "ASW" == data[item].get("Type") or "PSW" == data[item].get("Type"):
                            if __header__ == "ELOC Recheck With Tool" or __header__ == "LOC Complete" or __header__ == "OPL/Defect":
                                key = ""
                                if __header__ == "ELOC Recheck With Tool":
                                    key = "Assigned task (ELOC)"
                                elif __header__ == "LOC Complete":
                                    key = "Delivered task (ELOC)"
                                elif __header__ == "OPL/Defect":
                                    key = "Nums of Defect"
                                else:
                                    logger.error("BUG")
                                
                                if __header__ == "ELOC Recheck With Tool" or __header__ == "LOC Complete":
                                    data_prj['data'][key][data[item].get("Type")] += float(data[item].get(__header__))
                                elif __header__ == "OPL/Defect":
                                    if "Defect" == data[item].get("OPL/Defect"):
                                        data_prj['data'][key][data[item].get("Type")] += 1
                                else:
                                    logger.error("BUG")
                            elif __header__ == "Planned Start" or __header__ == "Planned End" or __header__ == "Release Date":
                                if __header__ == "Planned Start":
                                    key = "Assigned date"
                                elif __header__ == "Planned End":
                                    key = "Target date"
                                elif __header__ == "Release Date":
                                    key = "Delivered date"
                                else:
                                    logger.error("BUG")
                                
                                data_prj['data'][key][data[item].get("Type")].append(data[item].get(__header__))
                            else:
                                logger.error("BUG")
                        else:
                            logger.error("BUG")

        for __header__ in lst_header:
            if __header__ == "Planned Start" or __header__ == "Planned End" or __header__ == "Release Date":
                if __header__ == "Planned Start":
                    key = "Assigned date"
                elif __header__ == "Planned End":
                    key = "Target date"
                elif __header__ == "Release Date":
                    key = "Delivered date"
                else:
                    logger.error("BUG")

                for __type__ in ["ASW", "PSW"]:
                    if (len(data_prj['data'][key][__type__]) == 0):
                        del data_prj['data'][key][__type__]
                        data_prj['data'][key][__type__] = "NA"
                    else:
                        previous_lst_date = list(data_prj['data'][key][__type__])
                        del data_prj['data'][key][__type__]
                        if key == "Assigned date":
                            data_prj['data'][key][__type__] = min(previous_lst_date)
                        else:
                            data_prj['data'][key][__type__] = max(previous_lst_date)

        for __type__ in ["ASW", "PSW"]:
            data_prj['data']['Remain (ELOC)'][__type__] = data_prj['data']['Assigned task (ELOC)'][__type__] - data_prj['data']['Delivered task (ELOC)'][__type__]

            if data_prj['data']['Assigned task (ELOC)'][__type__] > 0:
                data_prj['data']['% Completion'][__type__] = round(data_prj['data']['Delivered task (ELOC)'][__type__]/data_prj['data']['Assigned task (ELOC)'][__type__] * 100,2)
            elif data_prj['data']['Assigned task (ELOC)'][__type__] == 0:
                data_prj["Type"].remove(__type__)
                for key, val in dict(data_prj['data']).items():
                    del data_prj['data'][key][__type__]
            else:
                data_prj['data']['% Completion'][__type__] = "NG"

        
        result[index] = dict(data_prj)

    print_information_deliverables(result, header=result_header)

    return result

def print_information_deliverables(data, header):
    print(','.join(header))
    
    result = dict()

    flag = False
    for index, obj_lv1 in dict(data).items():
        for __type__ in obj_lv1["Type"]:
            str_output = list()
            str_output.append(obj_lv1["Project"])
            str_output.append(__type__)
            for __header__ in header:
                for key, value in dict(obj_lv1["data"]).items():
                    if key == __header__:
                        if __type__ in list(value.keys()):
                            flag = True
                            str_output.append(str(value[__type__]))
                        else:
                            flag = False
                            break
                if(flag == False):
                    continue
            
            if(flag):
                print(','.join(str_output))


def make_folder_release(path_summary, l_packages, dir_output, begin=59, end=59):
    logger.debug("make_folder_release")
    try:
        doc = FileSummaryXLSX(path_summary)
        data = doc.parse2json(begin=begin, end=end)

        print("Start Making Folder Release")
        for package in l_packages:
            data_package = doc.get_data(data=data, key="Package", value=package)
            path_package = Path(dir_output).joinpath(str(package))
            Path(path_package).mkdir(exist_ok=True)
            if utils.load(CONST.SETTING).get("sheetname") == "Merged_COEM":
                path_package = Path(path_package).joinpath(str("COEM"))
            else:
                path_package = Path(path_package).joinpath(str("JOEM"))
            Path(path_package).mkdir(exist_ok=True)

            for item in data_package.keys():
                taskid = data_package[item].get("TaskID")
                if taskid is not "None":
                    date_end = data_package[item].get("Planned End") if data_package[item].get("Planned End") != "None" else None
                    if date_end is not None:
                        date_end = datetime.datetime.strptime(date_end, "%Y-%m-%d %H:%M:%S").strftime("%d-%b-%Y")
                        dir_taskid = Path(path_package).joinpath(str(date_end), str(taskid))
                        dir_taskid_RV = Path(dir_taskid).joinpath(str("RV"))
                        dir_taskid_AR = Path(dir_taskid).joinpath(str("AR"))
                        Path(dir_taskid).parent.mkdir(parents=True, exist_ok=True)
                        Path(dir_taskid).mkdir(exist_ok=True)
                        Path(dir_taskid_RV).mkdir(exist_ok=True)
                        Path(dir_taskid_AR).mkdir(exist_ok=True)

                        print("{},{},{}".format(package, taskid,"OK"))
                    else:
                        logger.warning("Not filling Planned Start/Planned End: {},{},{}".format(package, taskid,"OK"))

        print("Finish Making Folder Release")
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done")

# Main
def main():
    try:
        print("<<Checker Version: {}>>".format(utils.load(CONST.VERSION).get("version")))

        file_summary = utils.load(CONST.SETTING).get("file_summary")
        sheetname = utils.load(CONST.SETTING).get("sheetname")
        dir_output = utils.load(CONST.SETTING).get("dir_output")

        l_taskids = utils.load(CONST.SETTING).get("l_taskids_coem")
        dir_input = utils.load(CONST.SETTING).get("dir_input_coem")

        lst_opt = utils.load(CONST.SETTING).get("mode_coem")

        if sheetname == "Merged_JOEM":
            l_taskids = utils.load(CONST.SETTING).get("l_taskids_joem")
            dir_input = utils.load(CONST.SETTING).get("dir_input_joem")
            lst_opt = utils.load(CONST.SETTING).get("mode_joem")

        if(not(Path(file_summary).exists() \
            and sheetname is not None \
            and Path(dir_output).exists() \
            and Path(dir_input).exists() \
            and len(l_taskids)) \
        ):
            logger.error("Please make sure file_summary, sheetname, dir_input, dir_output, l_taskids are exist!\n\t+ file_summary: {}\n\t+ sheetname: {}\n\t+ dir_output: {}\n\t+ dir_input: {}\n\t+ l_taskids: {}".format(file_summary, sheetname, dir_output, dir_input, l_taskids))
            os.system("pause")
            quit()

        index_begin = utils.load(CONST.SETTING).get("coordinator").get("begin")
        index_end = utils.load(CONST.SETTING).get("coordinator").get("end")

        for opt in lst_opt:
            if opt == "check_releases":
                check_releases(path_summary=file_summary, dir_input=dir_input, taskids=l_taskids, begin=index_begin, end=index_end)
            elif opt == "check_archives":
                if sheetname == "Merged_JOEM":
                    check_archives_joem(path_summary=file_summary, dir_input=dir_input, taskids=l_taskids, begin=index_begin, end=index_end)
                else:
                    check_archives(path_summary=file_summary, dir_input=dir_input, taskids=l_taskids, begin=index_begin, end=index_end)
            elif opt == "make_archives":
                if sheetname == "Merged_JOEM":
                    make_archives_joem(path_summary=file_summary, dir_input=dir_input, dir_output=dir_output, taskids=l_taskids, begin=index_begin, end=index_end)
                else:
                    make_archieves(path_summary=file_summary, dir_input=dir_input, dir_output=dir_output, taskids=l_taskids, begin=index_begin, end=index_end)
            elif opt == "make_folder_release":
                """Make folder release"""
                l_folder_package = utils.load(CONST.SETTING).get("l_folder_package")
                make_folder_release(path_summary=file_summary, l_packages=l_folder_package, dir_output=dir_output, begin=index_begin, end=index_end)
            elif opt == "create_summary_json_file":
                """Create json file of summary to backup"""
                create_summary_json_file(file_summary=file_summary, sheetname="Merged_COEM", begin=index_begin, end=index_end)
                create_summary_json_file(file_summary=file_summary, sheetname="Merged_JOEM", begin=index_begin, end=index_end)
            elif opt == "collect_information_deliverables":
                """Collect information for deliverables"""
                collect_information_deliverables(file_summary=file_summary, sheetname="Merged_COEM", begin=index_begin, end=index_end)
                collect_information_deliverables(file_summary=file_summary, sheetname="Merged_JOEM", begin=index_begin, end=index_end)
            else:
                logger.error("I dont know your mode")
    except Exception as e:
        logger.exception(e)
    finally:
        logger.debug("Done)")

def test():
    directory = "C:/Users/nhi5hc/Desktop/Input/Test/MT_611_CRB_BoundByRegenProhibit_Gradient_SetValue"

    print(Path(utils.scan_files(Path(directory).as_posix(), ext='.xlsm')[0][1]).as_posix())
    lst_file_test_design = [f for f in utils.scan_files(Path(directory).as_posix(), ext='.xlsm')[0] if re.search('Documents/TD_.*.xlsm', str(Path(f).as_posix))]

def check_update_version():
    flag = False
    try:
        directory="//hc-ut40070c/duongnguyen/9000_utils/hieu.nguyen-trung/script_auto_checker"
        data = utils.scan_files(directory, ext="version.json")[0]

        current_version = utils.load(CONST.VERSION).get("version")

        l_version = []
        for f in data:
            version = utils.load(f).get("version")
            l_version.append(version)

        l_version = list(set(l_version))
        l_version.sort()
        latest_version = ""
        if (len(l_version) > 0):
            latest_version = l_version[-1]
        else:
            logger.error("BUG check update version")

        flag = False

        print("-----------------------------------------------------------------")
        if latest_version != None and latest_version != current_version:
            print('Please checkout new version "{}" at: "{}"'.format(latest_version, directory))
            flag = True
        else:
            print("Your version is latest")
        print("-----------------------------------------------------------------")

    except Exception as e:
        flag = False
        logger.exception(e)
    finally:
        logger.debug("Done)")
        return flag

if __name__ == "__main__":
    if(check_update_version()):
        pass
        # print("If you do not upgrade new version, you have to wait in 10s")
        # for i in range(0, 11):
        #     print("------Waiting: {} s-------".format(str(9 - i)))
        #     time.sleep(1)

    # test()
    main()
    os.system("pause")
